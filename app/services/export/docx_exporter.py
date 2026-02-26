"""
services/export/docx_exporter.py — Экспорт синопсиса в .docx по шаблону.

Стратегия:
1. Для пустых ячеек (Row 0,1,2,3,4,6,7,24,28) → вставляем текст целиком
2. Для ячеек с шаблонным текстом (Row 8,9,11,12...) → заменяем пробелы-плейсхолдеры
   на конкретные значения (названия препаратов, числа)
3. Для ячеек с полностью кастомным текстом (Row 13,14) → заменяем содержимое
"""

import copy
import math
import re
import os
from typing import Any, Dict, Optional, List
from docx import Document
from docx.shared import Pt

try:
    from app.utils.study_timeline import calculate_timeline
except ImportError:
    try:
        from study_timeline import calculate_timeline
    except ImportError:
        calculate_timeline = None


def _replace_blanks_in_runs(cell, replacements: list):
    """
    Заменяет пробелы-плейсхолдеры в runs ячейки.
    
    Шаблон использует два вида плейсхолдеров:
    1. Серии EN SPACE (\u2002) — Unicode пробелов шириной n-space
    2. Пустые runs ('')  подряд перед/после EN spaces
    3. Серии обычных пробелов (' {5,}')
    4. 'ХХ' как плейсхолдер
    
    Стратегия: для каждого параграфа находим «группы пустых runs»
    (runs содержащие только \u2002 или '') и заменяем первый run 
    группы на значение, остальные очищаем.
    """
    idx = 0
    BLANK_CHARS = {'\u2002', '\u2003', '\u2004', '\u2005', '\u2006', 
                   '\u2007', '\u2008', '\u2009', '\u200a', '\u00a0'}
    
    def _is_blank_run(text):
        """Run считается пустым/плейсхолдером."""
        if not text:
            return True
        return all(ch in BLANK_CHARS or ch == ' ' for ch in text)
    
    for para in cell.paragraphs:
        runs = para.runs
        i = 0
        while i < len(runs) and idx < len(replacements):
            run = runs[i]
            
            # Вариант 1: серия EN spaces в одном run
            if re.search(r'[\u2002\u2003\u00a0]{3,}', run.text):
                run.text = re.sub(r'[\u2002\u2003\u00a0]{3,}', 
                                  replacements[idx], run.text, count=1)
                idx += 1
                i += 1
                continue
            
            # Вариант 2: серия обычных пробелов
            if re.search(r' {5,}', run.text):
                run.text = re.sub(r' {5,}', replacements[idx], run.text, count=1)
                idx += 1
                i += 1
                continue
            
            # Вариант 3: группа пустых runs подряд (пустые + EN spaces)
            if _is_blank_run(run.text):
                # Считаем сколько пустых runs подряд
                group_start = i
                while i < len(runs) and _is_blank_run(runs[i].text):
                    i += 1
                group_len = i - group_start
                
                # Если группа ≥ 3 runs — это плейсхолдер
                if group_len >= 3 and idx < len(replacements):
                    runs[group_start].text = replacements[idx]
                    for j in range(group_start + 1, group_start + group_len):
                        runs[j].text = ""
                    # Проверяем: если следующий run начинается с буквы/цифры,
                    # добавляем пробел в конец замены (иначе слипнется)
                    if i < len(runs) and runs[i].text and runs[i].text[0].isalnum():
                        runs[group_start].text += " "
                    idx += 1
                continue
            
            # Вариант 4: ХХ
            if 'ХХ' in run.text and idx < len(replacements):
                run.text = run.text.replace('ХХ', replacements[idx], 1)
                idx += 1
                i += 1
                continue
            
            i += 1


def _set_cell_text(cell, text: str):
    """
    Заменяет содержимое ячейки целиком на новый текст.
    Сохраняет шрифт из первого run.
    """
    # Найдём оригинальное форматирование run
    source_rpr = None
    for para in cell.paragraphs:
        for run in para.runs:
            rpr = run._element.find(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr'
            )
            if rpr is not None:
                source_rpr = copy.deepcopy(rpr)
                break
        if source_rpr is not None:
            break

    # Удаляем все параграфы кроме первого
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    # Очищаем первый параграф
    first_para = cell.paragraphs[0]
    for run in first_para.runs:
        run._element.getparent().remove(run._element)

    # Вставляем новый текст — каждая строка в новый параграф
    lines = [l.strip() for l in text.split('\n') if l.strip()]

    for i, line in enumerate(lines):
        if i == 0:
            para = first_para
        else:
            para = cell.add_paragraph()
            # Копируем стиль параграфа
            if first_para.style:
                para.style = first_para.style

        run = para.add_run(line)

        # Применяем оригинальное форматирование
        if source_rpr is not None:
            clean_rpr = copy.deepcopy(source_rpr)
            # Убираем highlight (жёлтое выделение из шаблона)
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            for highlight_el in clean_rpr.findall(f'{ns}highlight'):
                clean_rpr.remove(highlight_el)
            # Также убираем shd (background shading)
            for shd_el in clean_rpr.findall(f'{ns}shd'):
                clean_rpr.remove(shd_el)
            run._element.insert(0, clean_rpr)
        else:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)


def _num_to_text(n: int) -> str:
    """Число → текст (для 'четырёх периодов')."""
    mapping = {
        1: "одного", 2: "двух", 3: "трёх", 4: "четырёх",
        5: "пяти", 6: "шести",
    }
    return mapping.get(n, str(n))


def export_synopsis(
    pipeline_result: Dict[str, Any],
    template_path: str,
    output_path: str = "synopsis_output.docx",
) -> str:
    """
    Заполняет шаблон синопсиса данными из Pipeline.

    Args:
        pipeline_result: dict с ключами pk, design, sample_size, synopsis
        template_path: путь к шаблону .docx
        output_path: путь к выходному файлу

    Returns:
        путь к сохранённому файлу
    """
    doc = Document(template_path)
    table = doc.tables[0]

    # ── Распаковываем данные ──
    syn = pipeline_result.get("synopsis", {})

    pk = pipeline_result.get("pk", {})
    if hasattr(pk, "model_dump"):
        pk = pk.model_dump()

    design = pipeline_result.get("design", {})
    if hasattr(design, "model_dump"):
        design = design.model_dump()

    sample = pipeline_result.get("sample_size", {})
    if hasattr(sample, "model_dump"):
        sample = sample.model_dump()

    # Удобные переменные
    inn = syn.get("inn", "")
    test_drug = syn.get("test_drug", inn)
    ref_drug = syn.get("reference_drug", "________")
    dosage_form = syn.get("dosage_form", "") or syn.get("dosage_form_input", "") or pipeline_result.get("dosage_form", "")
    dosage = syn.get("dosage", "") or syn.get("dosage_input", "") or pipeline_result.get("dosage", "")
    n_total = str(sample.get("n_total", ""))
    n_per_group = str(int(sample.get("n_total", 0)) // 2) if sample.get("n_total") else ""
    washout = str(design.get("washout_days", "")) if design.get("washout_days") else ""
    t_half_val = ""
    if pk.get("t_half") and isinstance(pk["t_half"], dict):
        t_half_val = str(pk["t_half"].get("value", ""))
    n_blood = str(design.get("n_blood_points", ""))
    sampling_h = str(design.get("sampling_duration_hours", ""))

    # Критерии включения/невключения: параметры из pipeline_result
    # sex_restriction может быть в syn (от LLM) или в pipeline_result напрямую
    sex_restriction = (
        syn.get("sex_restriction")
        or pipeline_result.get("sex_restriction", "males_only")
    )
    # Приведём к строке для criteria_generator
    if hasattr(sex_restriction, "value"):
        sex_restriction = sex_restriction.value
    if sex_restriction in ("М", "M", "males_only", "male"):
        sex_restriction = "males_only"
    elif sex_restriction in ("МЖ", "males_and_females", "both"):
        sex_restriction = "males_and_females"
    
    age_max = syn.get("age_max") or pipeline_result.get("age_max", 45)
    age_min = syn.get("age_min") or pipeline_result.get("age_min", 18)

    # T½ в часах для расчёта контрацепции
    t_half_hours = None
    if pk.get("t_half_hours"):
        t_half_hours = pk["t_half_hours"]
    elif pk.get("t_half") and isinstance(pk["t_half"], dict):
        t_half_hours = pk["t_half"].get("value")

    # ════════════════════════════════════════════
    # ПУСТЫЕ ЯЧЕЙКИ — вставляем текст целиком
    # ════════════════════════════════════════════

    # Row 0: Название протокола
    if syn.get("protocol_title"):
        _set_cell_text(table.rows[0].cells[1], syn["protocol_title"])

    # Row 1: Идентификационный номер протокола
    study_id = syn.get("study_id") or pipeline_result.get("study_id", "")
    if study_id:
        _set_cell_text(table.rows[1].cells[1], study_id)

    # Row 2: Спонсор
    sponsor_text = (
        syn.get("sponsor")
        or pipeline_result.get("sponsor_name")
        or pipeline_result.get("sponsor", "")
    )
    if sponsor_text and sponsor_text != "________":
        _set_cell_text(table.rows[2].cells[1], sponsor_text)

    # Row 3: Исследовательский центр
    center_text = (
        syn.get("research_center")
        or pipeline_result.get("research_center", "")
    )
    if center_text and center_text != "________":
        _set_cell_text(table.rows[3].cells[1], center_text)

    # Row 4: Биоаналитическая лаборатория
    lab_text = (
        syn.get("bioanalytical_lab")
        or pipeline_result.get("bioanalytical_lab", "")
    )
    if lab_text and lab_text != "________":
        _set_cell_text(table.rows[4].cells[1], lab_text)

    # Row 6: Название исследуемого препарата
    _set_cell_text(table.rows[6].cells[1], test_drug)

    # Row 7: Действующее вещество
    _set_cell_text(table.rows[7].cells[1], inn)

    # ════════════════════════════════════════════
    # ЯЧЕЙКИ С ШАБЛОННЫМ ТЕКСТОМ — замена плейсхолдеров
    # Если шаблон пустой (нет runs) → вставляем текст целиком
    # ════════════════════════════════════════════

    # Определяем режим приёма
    # Приоритет: synopsis (из инструкции к препарату) > design (из Design Agent)
    intake_mode = syn.get("intake_mode") or design.get("intake_mode", "fasting")
    if isinstance(intake_mode, str):
        intake_text_map = {
            "fasting": "натощак",
            "fed": "после приёма высококалорийной пищи",
            "both": "натощак и после приёма высококалорийной пищи",
        }
        correct_intake = intake_text_map.get(intake_mode, intake_mode)
    else:
        correct_intake = "натощак"

    def _cell_has_content(cell) -> bool:
        """Проверяет есть ли контент (runs) в ячейке."""
        return any(len(p.runs) > 0 for p in cell.paragraphs)

    # Кратность приёма определяется дизайном (для Row 8, 9)
    _release = syn.get("release_type", "immediate")
    if _release == "modified" or design.get("intake_mode") == "both":
        dosing_word = "многократного"
    else:
        dosing_word = "однократного"

    # Row 8: Цель исследования
    cell8 = table.rows[8].cells[1]
    if _cell_has_content(cell8):
        # Заполненный шаблон — заменяем плейсхолдеры
        _replace_blanks_in_runs(cell8, [
            test_drug, ref_drug,
            test_drug, ref_drug,
        ])
        # Замена «натощак/после приема...»
        for para in cell8.paragraphs:
            for run in para.runs:
                if "натощак/после приема высококалорийной пищи" in run.text:
                    run.text = run.text.replace(
                        "натощак/после приема высококалорийной пищи", correct_intake)
                elif "натощак/после" in run.text:
                    run.text = run.text.replace("натощак/после", correct_intake)
                elif "натощак" in run.text and correct_intake != "натощак":
                    run.text = run.text.replace("натощак", correct_intake)
    else:
        # Пустой шаблон — вставляем полный текст
        objective_text = (
            f"Основная цель:\n"
            f"Оценка сравнительной фармакокинетики и биоэквивалентности "
            f"препаратов {test_drug} и референтного препарата {ref_drug} "
            f"{correct_intake} у здоровых добровольцев.\n"
            f"\n"
            f"Дополнительная цель:\n"
            f"Сравнительная оценка безопасности и переносимости {dosing_word} приема "
            f"препарата {test_drug} у здоровых добровольцев."
        )
        _set_cell_text(cell8, objective_text)

    # Row 9: Задачи исследования
    cell9 = table.rows[9].cells[1]
    if _cell_has_content(cell9):
        _replace_blanks_in_runs(cell9, [
            inn, test_drug + " и " + ref_drug,
        ])
    else:
        tasks_text = (
            f"Определить концентрацию {inn} в плазме крови добровольцев "
            f"после {dosing_word} применения сравниваемых препаратов "
            f"{test_drug} и {ref_drug}.\n"
            f"Оценить фармакокинетические параметры и относительную "
            f"биодоступность сравниваемых препаратов.\n"
            f"Оценить биоэквивалентность сравниваемых препаратов на основе "
            f"статистического анализа фармакокинетических параметров."
        )
        _set_cell_text(cell9, tasks_text)

    # Row 10: Дизайн исследования
    cell10 = table.rows[10].cells[1]
    if _cell_has_content(cell10):
        _replace_blanks_in_runs(cell10, [test_drug, ref_drug])
    else:
        # ── Дизайн исследования по спецификации ──
        d_type = design.get("design_type", "2x2_crossover")
        if hasattr(d_type, "value"):
            d_type = d_type.value
        n_periods_val = int(design.get("n_periods", 2) or 2)
        n_sequences_val = int(design.get("n_sequences", 2) or 2)
        washout_d = washout or "7"

        # Производители для форматирования названий
        _test_mfr_name = syn.get("manufacturer", "") or ""
        _test_mfr_country = syn.get("manufacturer_country", "") or ""
        _ref_mfr = syn.get("ref_manufacturer", "") or ""

        # Форматируем названия: Название, лек.форма, дозировка (Производитель, Страна)
        if _test_mfr_name:
            _mfr_str = f"{_test_mfr_name}, {_test_mfr_country}" if _test_mfr_country else _test_mfr_name
            _test_full = f"{test_drug} ({_mfr_str})"
        else:
            _test_full = test_drug
        if _ref_mfr:
            _ref_full = f"{ref_drug}, {dosage_form}, {dosage} ({_ref_mfr})"
        else:
            _ref_full = ref_drug

        # Название дизайна по спецификации
        _design_name_map = {
            # Простой перекрестный
            ("2x2_crossover", 2, 2): "двухпериодного перекрестного исследования в двух группах",
            ("3x3_crossover", 3, 3): "трехпериодного перекрестного исследования в трех группах",
            # Частично репликативный (полуповторный): 2 опции × 3 последовательности × 3 периода
            ("partial_replicate", 3, 3): "трехпериодного перекрестного исследования в трех группах с частично репликативной схемой",
            # Обратная совместимость для replicate_3_period
            ("replicate_3_period", 3, 3): "трехпериодного перекрестного исследования в трех группах с частично репликативной схемой",
            # Полный репликативный (повторный)
            ("full_replicate_3_period", 3, 2): "трехпериодного перекрестного исследования в двух группах с полной репликативной схемой",
            ("full_replicate_4_period", 4, 2): "четырехпериодного перекрестного исследования в двух группах с полной репликативной схемой",
            ("replicate_4_period", 4, 2): "четырехпериодного перекрестного исследования в двух группах с полной репликативной схемой",
            # Параллельный
            ("parallel_1_period", 1, 2): "однопериодного в двух параллельных группах исследования",
            ("parallel_2_period", 2, 2): "двухпериодного в двух параллельных группах исследования",
            ("parallel", 1, 2): "однопериодного в двух параллельных группах исследования",
            # Адаптивный (двухэтапный)
            ("adaptive_crossover", 2, 2): "двухэтапного двухпериодного перекрестного исследования в двух группах",
            ("adaptive_parallel_1", 1, 2): "двухэтапного однопериодного в двух параллельных группах исследования",
            ("adaptive_parallel_2", 2, 2): "двухэтапного двухпериодного в двух параллельных группах исследования",
        }
        _design_name = _design_name_map.get(
            (d_type, n_periods_val, n_sequences_val),
            None
        )
        if _design_name is None:
            # Fallback: формируем по шаблону
            _period_prefix = {1: "одно", 2: "двух", 3: "трех", 4: "четырех"}.get(n_periods_val, str(n_periods_val))
            _seq_text = {2: "двух", 3: "трех", 4: "четырех"}.get(n_sequences_val, str(n_sequences_val))
            _design_name = f"{_period_prefix}периодного перекрестного исследования в {_seq_text} группах"

        # Кратность дозы
        _release = syn.get("release_type", "immediate")
        if _release == "modified":
            _dose_text = "с приемом многократной дозы"
        else:
            _dose_text = "с приемом однократной дозы"

        # Основной текст
        design_text = (
            f"Согласно действующим Правилам Евразийского экономического союза (ЕАЭС) "
            f"для сравнительной оценки фармакокинетических параметров "
            f"исследуемого препарата {_test_full} "
            f"и препарата {_ref_full} "
            f"стандартным дизайном исследования биоэквивалентности является "
            f"проведение открытого рандомизированного {_design_name} "
            f"{_dose_text}."
        )

        # Отмывочный период (только если предусмотрен дизайном)
        _no_washout = d_type in ("parallel_1_period", "parallel", "adaptive_parallel_1")
        if not _no_washout and n_periods_val > 1:
            design_text += (
                f" При этом периоды приема препаратов исследования разделяются "
                f"отмывочным периодом продолжительностью не менее пяти периодов "
                f"полувыведения (Т1/2) действующего вещества."
            )

        _set_cell_text(cell10, design_text)

    # Row 11: Методология исследования
    cell11 = table.rows[11].cells[1]

    # Генерируем полный текст методологии
    n_periods_val = design.get("n_periods", 2)
    washout_val = int(design.get("washout_days", 7) or 7)
    n_total_int = int(n_total) if n_total else 24
    n_sequences_val = int(design.get("n_sequences", 2) or 2)
    seq_desc = design.get("sequences_description", "TR/RT") or "TR/RT"
    seq_desc = _normalize_sequences(seq_desc)
    follow_up_val = int(design.get("follow_up_days", 7) or 7)

    intake_mode = syn.get("intake_mode") or design.get("intake_mode", "fasting")
    if hasattr(intake_mode, "value"):
        intake_mode = intake_mode.value

    # T½ в часах
    t_half_hours = 0.0
    if pk.get("t_half"):
        th = pk["t_half"]
        if isinstance(th, dict):
            t_half_hours = float(th.get("value", 0))
        else:
            t_half_hours = float(th)

    # Tmax в часах
    tmax_hours_val = 1.0
    if pk.get("tmax"):
        tm = pk["tmax"]
        if isinstance(tm, dict):
            tmax_hours_val = float(tm.get("value", 1.0) or 1.0)
        elif isinstance(tm, (int, float)):
            tmax_hours_val = float(tm)

    # ── Расчёт точек забора крови ──
    # Данные приходят из Design Agent (study_design._plan_blood_sampling)
    # который уже вызвал blood_sampling.calculate_blood_sampling
    sampling_times_hours_list = design.get("sampling_times_hours", [])
    n_blood_val = int(design.get("n_blood_points", 18) or 18)
    sampling_hours_val = int(design.get("sampling_duration_hours", 24) or 24)

    # Если Design Agent не рассчитал (старая версия) — пробуем blood_sampling
    blood_result = None
    if not sampling_times_hours_list and t_half_hours > 0:
        try:
            from app.utils.blood_sampling import calculate_blood_sampling
        except ImportError:
            try:
                from blood_sampling import calculate_blood_sampling
            except ImportError:
                calculate_blood_sampling = None

        if calculate_blood_sampling:
            blood_result = calculate_blood_sampling(
                n_periods=n_periods_val,
                t_half_hours=t_half_hours,
                tmax_hours=tmax_hours_val,
                n_subjects=n_total_int,
            )
            n_blood_val = blood_result.fk_points_per_period
            sampling_hours_val = blood_result.sampling_duration_hours

    # Вычисляем timeline
    if calculate_timeline is not None:
        tl = calculate_timeline(
            n_periods=n_periods_val,
            washout_days=washout_val,
            sampling_hours=sampling_hours_val,
            n_blood_points=n_blood_val,
            n_total=n_total_int,
            n_sequences=n_sequences_val,
            sequences_description=seq_desc,
            follow_up_days=follow_up_val,
        )
        b = tl.blood
    else:
        tl = None
        b = None

    methodology_text, rand_table_data = _generate_methodology_text(
        tl=tl, b=b,
        n_periods=n_periods_val,
        washout_days=washout_val,
        follow_up_days=follow_up_val,
        sampling_hours=sampling_hours_val,
        n_total=n_total_int,
        n_sequences=n_sequences_val,
        seq_desc=seq_desc,
        inn=inn,
        t_half_hours=t_half_hours,
        intake_mode=intake_mode,
        test_drug=test_drug,
        ref_drug=ref_drug,
        blood_result=blood_result,
        sampling_times_list=sampling_times_hours_list,
    )

    if _cell_has_content(cell11):
        try:
            _replace_blanks_in_runs(cell11, _methodology_blanks_values(
                tl, b, n_periods_val, washout_val, follow_up_val,
                n_total_int, inn, t_half_hours,
            ))
        except (IndexError, TypeError):
            _set_cell_text(cell11, methodology_text)
    else:
        _set_cell_text(cell11, methodology_text)

    # Вставляем Word-таблицу рандомизации вместо плейсхолдера
    if rand_table_data:
        _insert_randomization_table(cell11, rand_table_data)

    # Row 12: Количество добровольцев (краткое — полный расчёт в Row 24)
    cell12 = table.rows[12].cells[1]
    cv_intra_val = sample.get("cv_intra_used", "")
    n_base = str(sample.get("n_base", "")) if sample.get("n_base") else ""
    n_with_dropout = str(sample.get("n_with_dropout", "")) if sample.get("n_with_dropout") else n_total
    dropout_pct = sample.get("dropout_rate", 0.15)
    screenfail_pct = sample.get("screenfail_rate", 0.15)

    _cell12_text = cell12.text.strip()
    if _cell12_text and len(_cell12_text) > 10:
        # Шаблон заполнен — подставляем значения в EN SPACE пробелы
        _replace_blanks_in_runs(cell12, [
            str(cv_intra_val) if cv_intra_val else "",
            str(cv_intra_val) if cv_intra_val else "",
            str(n_with_dropout),
            n_total,
        ])
    else:
        # Пустой — вставляем краткий текст с форматированием
        _insert_volunteers_count_formatted(
            cell12, inn=inn,
            cv_intra_val=cv_intra_val,
            n_with_dropout=str(n_with_dropout),
            n_total=n_total,
            dropout_pct=dropout_pct,
            screenfail_pct=screenfail_pct,
        )

    # Row 13: Критерии включения
    cell13 = table.rows[13].cells[1]
    if not _cell_has_content(cell13):
        # Blank template — генерируем критерии
        try:
            from criteria_generator import generate_inclusion_criteria
        except ImportError:
            from app.services.export.criteria_generator import generate_inclusion_criteria
        incl_text = syn.get("inclusion_criteria", "")
        if not incl_text or incl_text == "Критерии включения: требуется доработка.":
            incl_text = generate_inclusion_criteria(
                sex=sex_restriction,
                age_min=age_min,
                age_max=age_max,
                t_half_hours=t_half_hours,
            )
        _set_cell_text(cell13, incl_text)
    elif syn.get("inclusion_criteria") and syn["inclusion_criteria"] != "Критерии включения: требуется доработка.":
        _set_cell_text(cell13, syn["inclusion_criteria"])

    # Row 14: Критерии невключения
    cell14 = table.rows[14].cells[1]
    if not _cell_has_content(cell14):
        try:
            from criteria_generator import generate_non_inclusion_criteria
        except ImportError:
            from app.services.export.criteria_generator import generate_non_inclusion_criteria
        noninc_text = syn.get("exclusion_criteria", "")
        if not noninc_text or noninc_text == "Критерии невключения: требуется доработка.":
            noninc_text = generate_non_inclusion_criteria(
                sex=sex_restriction,
                inn_ru=inn,
            )
        _set_cell_text(cell14, noninc_text)
    elif syn.get("exclusion_criteria") and syn["exclusion_criteria"] != "Критерии невключения: требуется доработка.":
        _set_cell_text(cell14, syn["exclusion_criteria"])

    # Row 15: Критерии исключения
    cell15 = table.rows[15].cells[1]
    if not _cell_has_content(cell15):
        try:
            from criteria_generator import generate_exclusion_criteria
        except ImportError:
            from app.services.export.criteria_generator import generate_exclusion_criteria

        # Получаем Tmax из PK Agent результата (переменная pk, не syn)
        tmax_val = None
        if isinstance(pk, dict):
            tmax_obj = pk.get("tmax")
            if isinstance(tmax_obj, dict):
                tmax_val = tmax_obj.get("value")
            elif isinstance(tmax_obj, (int, float)):
                tmax_val = tmax_obj

        excl_text = generate_exclusion_criteria(
            sex=sex_restriction,
            tmax_hours=tmax_val,
            inn_ru=inn,
        )
        _set_cell_text(cell15, excl_text)

    # Row 16: Исследуемый препарат (T), доза, схема приёма
    cell16 = table.rows[16].cells[1]
    if not _cell_has_content(cell16):
        _set_cell_text(cell16, _build_drug_cell(
            drug_name=test_drug,
            inn=inn,
            dosage_form=dosage_form,
            dosage=dosage,
            design=design,
            n_total_int=int(n_total) if n_total else 24,
            syn=syn,
            is_reference=False,
            ref_drug=ref_drug,
        ))
    else:
        _replace_blanks_in_runs(cell16, [
            test_drug, dosage_form, dosage, inn,
        ])

    # Row 17: Референтный препарат (R), доза, схема приёма
    cell17 = table.rows[17].cells[1]
    if not _cell_has_content(cell17):
        _set_cell_text(cell17, _build_drug_cell(
            drug_name=ref_drug,
            inn=inn,
            dosage_form=dosage_form,
            dosage=dosage,
            design=design,
            n_total_int=int(n_total) if n_total else 24,
            syn=syn,
            is_reference=True,
            ref_drug=ref_drug,
        ))
    else:
        _replace_blanks_in_runs(cell17, [
            ref_drug, inn, dosage_form, dosage,
        ])

    # Row 18: Периоды исследования
    cell18 = table.rows[18].cells[1]
    if not _cell_has_content(cell18):
        periods_text = _generate_periods_text(
            tl=tl,
            n_periods=n_periods_val,
            washout_days=washout_val,
            follow_up_days=follow_up_val,
            sampling_hours=sampling_hours_val,
            intake_mode=intake_mode,
            sex_restriction=sex_restriction,
        )
        _set_cell_text(cell18, periods_text)

    # Row 19: Продолжительность исследования
    cell19 = table.rows[19].cells[1]
    if not _cell_has_content(cell19):
        duration_text = _generate_duration_text(
            tl=tl,
            n_periods=n_periods_val,
            washout_days=washout_val,
            follow_up_days=follow_up_val,
            sampling_hours=sampling_hours_val,
        )
        _set_cell_text(cell19, duration_text)

    # Row 20: Изучаемые фармакокинетические параметры
    cell20 = table.rows[20].cells[1]
    _cell20_text = cell20.text.strip()
    if _cell20_text and len(_cell20_text) > 5:
        # Шаблон заполнен — подставляем МНН в плейсхолдеры
        _replace_blanks_in_runs(cell20, [
            inn,            # фармакокинетика ____ по данным
            inn,            # ФК параметры ____
            inn,            # вторичные ____
        ])
    else:
        # Пустой — вставляем с форматированием (subscript, italic)
        _insert_pk_parameters_formatted(cell20, inn)

    # Row 21: Аналитический метод
    cell21 = table.rows[21].cells[1]
    _cell21_text = cell21.text.strip()
    if _cell21_text and len(_cell21_text) > 5:
        _replace_blanks_in_runs(cell21, [inn])
    else:
        _set_cell_text(cell21, (
            f"Для определения концентрации аналита в плазме крови здоровых "
            f"добровольцев будет применяться метод высокоэффективной жидкостной "
            f"хроматографии с тандемным масс-селективным детектированием "
            f"(ВЭЖХ-МС/МС). Полная валидация биоаналитического метода "
            f"определения {inn} в плазме крови будет проведена в соответствии "
            f"с рекомендациями ЕАЭС и стандартными процедурами аналитической "
            f"лаборатории."
        ))

    # Row 22: Критерии БЭ
    cell22 = table.rows[22].cells[1]
    _cell22_text = cell22.text.strip()
    be_lower = design.get("be_lower", 80.0) or 80.0
    be_upper = design.get("be_upper", 125.0) or 125.0
    if _cell22_text and len(_cell22_text) > 5:
        # Шаблон заполнен — подставляем МНН в пробел
        _replace_blanks_in_runs(cell22, [inn])
    else:
        _insert_be_criteria_formatted(cell22, inn, be_lower, be_upper)

    # Row 23: Анализ параметров безопасности
    cell23 = table.rows[23].cells[1]
    if not (cell23.text.strip() and len(cell23.text.strip()) > 5):
        safety_text = (
            "В исследовании оценка безопасности будет проводиться на основании "
            "оценки жизненно важных показателей, данных объективных исследований, "
            "лабораторных показателей и частоты регистрации и анализа "
            "нежелательных явлений (НЯ). Оценка НЯ будет проводиться "
            "на основании:\n"
            "Жалоб, изменений самочувствия;\n"
            "Результатов физикального осмотра;\n"
            "Результатов оценки основных показателей жизнедеятельности "
            "(АД, ЧСС, ЧДД, температура тела);\n"
            "Лабораторного мониторинга (клинический анализ крови, "
            "биохимический анализ крови, общий анализ мочи).\n"
            "\n"
            "Описание НЯ будет представлено согласно следующей схеме:\n"
            "дата и время возникновения НЯ (если применимо);\n"
            "описание НЯ;\n"
            "серьезность НЯ;\n"
            "тяжесть НЯ;\n"
            "предвиденность НЯ;\n"
            "наличие связи с препаратом исследования;\n"
            "предпринятые действия для лечения НЯ;\n"
            "предпринятые действия в отношении исследуемого препарата;\n"
            "сопутствующая терапия НЯ;\n"
            "исход/разрешение НЯ.\n"
            "\n"
            "Информация о НЯ/серьезных нежелательных явлениях (СНЯ) будет "
            "собираться с момента первого дозирования исследуемого препарата/"
            "препарата сравнения и далее в ходе всего исследования до визита "
            "завершения.\n"
            "Сведения о неблагоприятных с медицинской точки зрения событиях, "
            "зарегистрированных до приема препарата в рамках исследования, "
            "будут вынесены в отдельную группу. Данные события не будут учтены "
            "в статистическом анализе по безопасности и будут приведены "
            "в итоговом отчете только в виде справочной информации."
        )
        _set_cell_text(cell23, safety_text)

    # Row 24: Расчет размера выборки (пустая в шаблоне)
    if sample.get("calculation_description"):
        _set_cell_text(table.rows[24].cells[1], sample["calculation_description"])

    # Row 25: Методы статистического анализа
    cell25 = table.rows[25].cells[1]
    _cell25_text = cell25.text.strip()
    d_type = design.get("design_type", "2x2_crossover")
    if hasattr(d_type, "value"):
        d_type = d_type.value
    _anova_factors = _get_anova_factors(d_type)

    if _cell25_text and len(_cell25_text) > 20:
        # Шаблон заполнен — вставляем формулы в P2/P3 и подставляем blanks
        _insert_hypothesis_formulas(cell25)
        _replace_blanks_in_runs(cell25, [
            inn,             # после "Сmax ___" — МНН
            _anova_factors,  # после "факторы:" — список факторов ANOVA
            "",              # после "ЭКГ," — доп.обследования
        ])
    else:
        # Пустой — вставляем полный текст с форматированными формулами
        _insert_stat_methods_formatted(cell25, inn, _anova_factors)

    # Row 26: Заслепление и Рандомизация
    cell26 = table.rows[26].cells[1]
    _cell26_text = cell26.text.strip()

    # Параметры рандомизации
    d_type = design.get("design_type", "2x2_crossover")
    if hasattr(d_type, "value"):
        d_type = d_type.value
    n_sequences_val = int(design.get("n_sequences", 2) or 2)
    seq_desc = design.get("sequences_description", "TR/RT") or "TR/RT"
    seq_norm = _normalize_sequences(seq_desc)
    seqs = seq_norm.split("/") if seq_norm else ["TR", "RT"]
    _n_per_grp = str(int(n_total) // max(n_sequences_val, 1)) if n_total else ""

    if _cell26_text and len(_cell26_text) > 20:
        # Шаблон заполнен — подставляем n_per_group и последовательности
        # Заменяем "T/R или R/T" на актуальные последовательности
        _seq_display = "/".join(seqs)  # "TRTR/RTRT"
        for para in cell26.paragraphs:
            for run in para.runs:
                if "T/R или R/T" in run.text:
                    run.text = run.text.replace("T/R или R/T", _seq_display)
                elif "–T/R" in run.text:
                    run.text = run.text.replace("–T/R", "– " + _seq_display)
        _replace_blanks_in_runs(cell26, [_n_per_grp])
    else:
        # Пустой — генерируем полный текст
        blind_text = _generate_blinding_randomization_text(
            n_sequences=n_sequences_val,
            seqs=seqs,
            n_per_group=_n_per_grp,
            design_type=d_type,
        )
        _set_cell_text(cell26, blind_text)

    # Row 27: Этические и регуляторные аспекты
    cell27 = table.rows[27].cells[1]
    insurance_text = (
        syn.get("insurance_company")
        or pipeline_result.get("insurance_company", "")
    )
    # Если insurance_text содержит \n — это полные реквизиты (название + адрес + телефон)
    # Извлекаем название для вставки в текст, а полные реквизиты — отдельным блоком
    ins_name = insurance_text.split("\n")[0] if insurance_text else ""
    ins_details = insurance_text if insurance_text else ""

    _cell27_text = cell27.text.strip()
    if _cell27_text and len(_cell27_text) > 20:
        # Шаблон заполнен — подставляем страховую компанию
        _replace_blanks_in_runs(cell27, [
            ins_name if ins_name and ins_name != "________" else "_____",
        ])
    else:
        # Пустой — генерируем текст
        _ins = ins_name if ins_name and ins_name != "________" else "_____"
        # Полные реквизиты страховой (адрес, телефон) — если есть
        _ins_full_block = ""
        if ins_details and "\n" in ins_details and ins_details != "________":
            _ins_full_block = f"\n{ins_details}"
        elif ins_name and ins_name != "________":
            _ins_full_block = f"\n{ins_name}"

        ethics_text = (
            "Исследование будет проводиться согласно Протоколу, "
            "в строгом соответствии с:\n"
            "Конституцией Российской Федерации;\n"
            "Этическими принципами Хельсинкской декларации Всемирной "
            "медицинской ассоциации 1964 г. в последней редакции, принятой "
            "на 75-ой Генеральной Ассамблее ВМА, Хельсинки, Финляндия, "
            "октябрь 2024 г;\n"
            "Решением Совета Евразийской экономической комиссии от "
            "03.11.2016 № 79 «Об утверждении Правил Надлежащей клинической "
            "практики Евразийского экономического союза»;\n"
            "Решением Совета Евразийской экономической комиссии от "
            "03.11.2016 № 85 «Об утверждении Правил проведения исследований "
            "биоэквивалентности лекарственных препаратов в рамках "
            "Евразийского экономического союза»;\n"
            "а также в соответствии с применимыми требованиями "
            "законодательства Российской Федерации с действующими "
            "изменениями.\n"
            f"Страхование жизни и здоровья добровольцев осуществляется "
            f"компанией:"
            f"{_ins_full_block}"
        )
        _set_cell_text(cell27, ethics_text)

    # Row 28: Номер версии Протокола и дата
    _proto_version = (
        pipeline_result.get("protocol_version")
        or syn.get("protocol_version")
        or 1
    )
    from datetime import date as _date_cls
    _today = _date_cls.today().strftime("%d.%m.%Y")
    _version_text = f"Версия {_proto_version}.0 от {_today}"
    _set_cell_text(table.rows[28].cells[1], _version_text)

    # ════════════════════════════════════════════
    # СОХРАНЯЕМ
    # ════════════════════════════════════════════

    # ── Убираем все highlight/mark из документа ──
    _remove_all_highlights(doc)

    doc.save(output_path)
    return output_path


def _remove_all_highlights(doc: Document):
    """Убирает жёлтое выделение (highlight) из всех runs во всех таблицах и параграфах."""
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    def _clean_element(element):
        for rpr in element.iter(f'{ns}rPr'):
            for hl in rpr.findall(f'{ns}highlight'):
                rpr.remove(hl)
            for shd in rpr.findall(f'{ns}shd'):
                rpr.remove(shd)

    # Параграфы в теле документа
    for para in doc.paragraphs:
        _clean_element(para._element)

    # Таблицы
    for table in doc.tables:
        _clean_element(table._tbl)


# ════════════════════════════════════════════
# ГЕНЕРАТОР ТЕКСТА МЕТОДОЛОГИИ (Row 11)
# ════════════════════════════════════════════

def _fmt_days(days: list, prefix: str = "День") -> str:
    """'День 1 и День 9' или 'День 1, День 9, День 17 и День 25'."""
    items = [f"{prefix} {d}" if prefix else str(d) for d in days]
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} и {items[1]}"
    return ", ".join(items[:-1]) + " и " + items[-1]


def _periods_desc(n: int) -> str:
    """'двух периодов ФК исследования (Период 1 и Период 2 соответственно) с отмывочным периодом между ними'"""
    w = {2: "двух", 3: "трёх", 4: "четырёх"}.get(n, str(n))
    names = [f"Период {i}" for i in range(1, n + 1)]
    if n <= 2:
        names_str = " и ".join(names) + " соответственно"
        wo = "с отмывочным периодом между ними"
    else:
        names_str = ", ".join(names[:-1]) + " и " + names[-1]
        wo = "с отмывочными периодами между ними"
    return f"{w} периодов ФК исследования ({names_str}) {wo}"


def _periods_word_gen(n: int) -> str:
    return {2: "двух периодах", 3: "трёх периодах", 4: "четырёх периодах"}.get(n, f"{n} периодах")


def _build_drug_cell(
    drug_name: str, inn: str, dosage_form: str, dosage: str,
    design: dict, n_total_int: int, syn: dict,
    is_reference: bool, ref_drug: str = "",
) -> str:
    """
    Генерирует текст ячейки Row 16 (T) или Row 17 (R) по формату шаблона.

    Формат:
        [Название], МНН – [МНН]
        Лекарственная форма: [форма]
        Дозировка: [дозировка]
        Состав на 1 [таблетку]: ...
        Действующее вещество: [МНН]
        Вспомогательные вещества: ...
        Схема приема: Каждый доброволец примет однократно натощак по [1 таблетке]
          препарата [название] в последовательности, предусмотренной схемой
          рандомизации для каждой группы (в День X или День Y ...).
        Условия хранения: ...
        Производитель: ...
        [для R: блок про выбор референтного препарата]
    """
    blank = "_____"

    # Единица лек. формы
    fl = dosage_form.lower() if dosage_form else ""
    if "таблет" in fl:
        unit_form = "1 таблетке"
        unit_noun = "таблетку"
    elif "капсул" in fl:
        unit_form = "1 капсуле"
        unit_noun = "капсулу"
    else:
        unit_form = f"1 единице ({dosage_form or blank})"
        unit_noun = dosage_form or blank

    # Натощак/после еды
    # Приоритет: synopsis (из инструкции к препарату) > design (из Design Agent)
    intake_mode = syn.get("intake_mode") or design.get("intake_mode", "fasting")
    if hasattr(intake_mode, "value"):
        intake_mode = intake_mode.value
    intake_word = {
        "fasting": "натощак",
        "fed": "после приёма высококалорийной пищи",
        "both": "натощак и после приёма высококалорийной пищи",
    }.get(intake_mode, "натощак")

    # Дни приёма (сквозная нумерация)
    n_periods_val = int(design.get("n_periods", 2) or 2)
    washout_val = int(design.get("washout_days", 7) or 7)
    dosing_days = [1 + i * washout_val for i in range(n_periods_val)]
    days_str = _fmt_days(dosing_days)

    # Данные
    if is_reference:
        manufacturer = syn.get("ref_manufacturer", "") or blank
        excipients = syn.get("ref_excipients", "") or blank
        storage_cond = syn.get("ref_storage_conditions", "") or blank
        composition = syn.get("ref_composition", "") or ""
        ref_ru = syn.get("ref_ru_number", "") or blank
    else:
        manufacturer = syn.get("manufacturer", "") or blank
        excipients = syn.get("excipients", "") or blank
        storage_cond = syn.get("storage_conditions", "") or blank
        composition = syn.get("composition", "") or ""
        ref_ru = ""

    L = []

    # Заголовок
    L.append(f"{drug_name or blank}, МНН – {inn}")

    # Форма и дозировка
    L.append(f"Лекарственная форма: {dosage_form or blank}")
    L.append(f"Дозировка: {dosage or blank}")

    # Состав
    L.append(f"Состав на {unit_form}:")
    L.append(f"Действующее вещество: {inn}")
    L.append(f"Вспомогательные вещества: {excipients}")

    # Схема приёма
    L.append("")
    L.append(
        f"Схема приема: Каждый доброволец примет однократно {intake_word} "
        f"по {unit_form} препарата {drug_name or blank} в "
        f"последовательности, предусмотренной схемой рандомизации для "
        f"каждой группы (в {days_str} соответствующего периода ФК "
        f"исследования)."
    )

    # Хранение и производитель
    L.append(f"Условия хранения: {storage_cond}")
    L.append(f"Производитель: {manufacturer}")

    # Для R — блок про выбор референтного
    if is_reference:
        L.append("")
        L.append(
            f"Выбор референтного препарата осуществлялся в соответствии с "
            f"Правилами проведения исследований биоэквивалентности "
            f"лекарственных средств Евразийского Экономического Союза, "
            f"утвержденными Решением Совета Евразийской экономической "
            f"комиссии от 3 ноября 2016 г. № 85."
        )
        L.append(
            f"Препарат {drug_name or blank} является оригинальным "
            f"препаратом, официально зарегистрированным и разрешенным для "
            f"медицинского применения на территории РФ "
            f"(РУ ЛП-№{ref_ru})."
        )
        L.append(
            f"Сравниваемые препараты содержат одно и то же действующее "
            f"вещество, выпускаются в одинаковой лекарственной форме – "
            f"{dosage_form or blank}."
        )

    return "\n".join(L)


def _insert_randomization_table(cell, rand_data: dict):
    """
    Вставляет Word-таблицу рандомизации в ячейку,
    заменяя параграф с <<<RANDOMIZATION_TABLE>>>.

    Args:
        cell: ячейка docx Table
        rand_data: {"seqs": ["TRR","RTR","RRT"], "n_per_group": 25,
                     "n_periods": 3, "intake_word": "натощак"}
    """
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm
    from docx.oxml import OxmlElement

    seqs = rand_data["seqs"]
    n_per_group = rand_data["n_per_group"]
    n_periods = rand_data["n_periods"]
    intake_word = rand_data["intake_word"]

    # Найдём параграф с плейсхолдером
    placeholder_para = None
    placeholder_idx = None
    for idx, para in enumerate(cell.paragraphs):
        if "<<<RANDOMIZATION_TABLE>>>" in para.text:
            placeholder_para = para
            placeholder_idx = idx
            break

    if placeholder_para is None:
        return

    # Сохраняем форматирование из первого run ячейки
    source_rpr = None
    for p in cell.paragraphs:
        for r in p.runs:
            rpr = r._element.find(qn('w:rPr'))
            if rpr is not None:
                source_rpr = copy.deepcopy(rpr)
                break
        if source_rpr is not None:
            break

    # Создаём таблицу: (len(seqs)+1) rows × (n_periods+1) cols
    n_rows = len(seqs) + 1  # header + groups
    n_cols = n_periods + 1    # label + periods

    # Создаём XML таблицы
    tbl_elem = OxmlElement('w:tbl')

    # Свойства таблицы
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    # Borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tbl_elem.append(tblPr)

    # Grid
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(n_cols):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), '2000')
        tblGrid.append(gridCol)
    tbl_elem.append(tblGrid)

    def _make_cell(text, bold=False):
        tc = OxmlElement('w:tc')
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '18')  # 9pt
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '18')
        rPr.append(szCs)
        if bold:
            b_elem = OxmlElement('w:b')
            rPr.append(b_elem)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    # Header row
    header_row = OxmlElement('w:tr')
    header_row.append(_make_cell("", bold=True))
    for i in range(1, n_periods + 1):
        header_row.append(_make_cell(f"Период {i}", bold=True))
    tbl_elem.append(header_row)

    # Data rows
    for g_idx, seq in enumerate(seqs):
        tr = OxmlElement('w:tr')
        tr.append(_make_cell(f"Группа {g_idx+1} (n={n_per_group}) ({seq})", bold=True))
        for ch in seq:
            label = "[T]" if ch.upper() == "T" else "[R]"
            tr.append(_make_cell(f"однократно {intake_word} {label}"))
        tbl_elem.append(tr)

    # Вставляем таблицу вместо плейсхолдера
    placeholder_elem = placeholder_para._element
    placeholder_elem.getparent().replace(placeholder_elem, tbl_elem)


def _normalize_sequences(raw: str) -> str:
    """Нормализует sequences_description в формат TRTR/RTRT."""
    if not raw:
        return "TR/RT"
    if re.match(r'^[TR/]+$', raw.strip()):
        return raw.strip()
    sequences = []
    parts = re.split(r'Последовательность\s*\d+\s*[:\s]+', raw, flags=re.IGNORECASE)
    for part in parts:
        part = part.strip().rstrip('.')
        if not part:
            continue
        letters = re.findall(r'[TR]', part.upper())
        if letters:
            sequences.append("".join(letters))
    return "/".join(sequences) if sequences else "TR/RT"


def _generate_methodology_text(
    tl, b,
    n_periods, washout_days, follow_up_days, sampling_hours,
    n_total, n_sequences, seq_desc, inn, t_half_hours,
    intake_mode, test_drug="", ref_drug="",
    blood_result=None,
    sampling_times_list=None,
) -> str:
    """Генерирует полный текст раздела «Методология исследования»."""

    n_per_group = n_total // max(n_sequences, 1)

    # Intake
    if intake_mode == "fed":
        intake_word = "после приёма высококалорийной пищи"
        intake_cond = ("после приёма стандартного высококалорийного завтрака, "
                       "запивая 200 мл ±10 мл бутилированной негазированной "
                       "воды комнатной температуры")
    else:
        intake_word = "натощак"
        intake_cond = ("натощак (не менее чем через 8 часов после последнего "
                       "приема пищи), запивая 200 мл ±10 мл бутилированной "
                       "негазированной воды комнатной температуры, которая "
                       "должна быть выпита полностью")

    # T½ format
    if t_half_hours and t_half_hours > 0:
        t_half_str = f"{t_half_hours:.2f}" if t_half_hours < 1 else f"{t_half_hours:.1f}"
    else:
        t_half_str = "_____"

    # Timeline data
    if tl is not None:
        fk_period_days = tl.fk_period_days
        dosing_days = tl.dosing_days
        hosp_days = [p.hospitalization_day for p in tl.fk_periods]
        discharge_days = [p.sampling_end_day for p in tl.fk_periods]
        fk_pts = b.fk_points_per_period if b else 18
        fk_total = b.fk_total_points if b else fk_pts * n_periods
        flush_count = fk_total - n_periods
        flush_ml = b.flush_volume_ml if b else flush_count * 0.5
        fk_total_ml = b.fk_total_ml if b else fk_total * 5 + flush_ml
        lab_ml = b.lab_total_ml if b else 66
        total_ml = b.total_per_subject_ml if b else fk_total_ml + lab_ml
        biosamples = b.biosamples_total if b else fk_pts * n_periods * n_total
    else:
        fk_period_days = 2
        dosing_days = [1 + i * washout_days for i in range(n_periods)]
        hosp_days = [d - 1 for d in dosing_days]
        discharge_days = [d + 1 for d in dosing_days]
        fk_pts = 18
        fk_total = fk_pts * n_periods
        flush_count = fk_total - n_periods
        flush_ml = flush_count * 0.5
        fk_total_ml = fk_total * 5 + flush_ml
        lab_ml = (2 + n_periods) * 14 + 10
        total_ml = fk_total_ml + lab_ml
        biosamples = fk_pts * n_periods * n_total

    last_dose = dosing_days[-1]
    followup_day = last_dose + follow_up_days
    last_period = f"Периоде {n_periods}"
    period_names = [f"Период {i}" for i in range(1, n_periods + 1)]
    period_names_str = ", ".join(period_names[:-1]) + " и " + period_names[-1] if len(period_names) > 2 else " и ".join(period_names)

    # Таблица рандомизации — данные для Word-таблицы
    seq_norm = _normalize_sequences(seq_desc)
    seqs = seq_norm.split("/") if seq_norm else ["TR", "RT"]
    # Сохраняем данные как структуру для вставки Word-таблицы
    _rand_table_data = {
        "seqs": seqs,
        "n_per_group": n_per_group,
        "n_periods": n_periods,
        "intake_word": intake_word,
    }

    L = []

    # Блок 1: Введение
    L.append(
        f"Настоящее исследование будет выполнено с участием здоровых "
        f"добровольцев, соответствующих критериям включения/невключения "
        f"и подписавших «Информационный листок добровольца с формой "
        f"информированного согласия». Исследование будет состоять из "
        f"следующих периодов: периода скрининга, "
        f"{_periods_desc(n_periods)} и периода последующего наблюдения."
    )

    # Блок 2: Длительности
    L.append(
        f"Длительность периода скрининга составит не более 14 дней, "
        f"длительность каждого периода ФК исследования составит "
        f"{fk_period_days} дней, длительность отмывочного периода – "
        f"{washout_days} дней от приема препарата в Периоде 1 ФК "
        f"исследования, периода последующего наблюдения – {follow_up_days} "
        f"дней от приема препарата в {last_period} ФК исследования."
    )

    # Блок 3: Рандомизация
    n_groups_word = {2: "двух", 3: "трёх", 4: "четырёх"}.get(len(seqs), str(len(seqs)))
    L.append(
        f"Добровольцы будут распределены в соответствии с "
        f"рандомизационным списком в одну из {n_groups_word} групп "
        f"в соотношении {'1:1' if len(seqs) == 2 else '1:' * (len(seqs)-1) + '1'}:"
    )
    L.append("<<<RANDOMIZATION_TABLE>>>")

    # Блок 4: Скрининг
    L.append(
        f"Период скрининга: после подписания Информационного листка "
        f"с формой информированного согласия, добровольцы пройдут "
        f"процедуры скрининга для оценки соответствия критериям "
        f"включения/невключения."
    )
    L.append(
        f"Дата подписания добровольцем формы информированного согласия "
        f"будет считаться датой его включения в клиническое исследование."
    )

    # Блок 5: Периоды ФК
    L.append("Периоды фармакокинетического (ФК) исследования")
    L.append(
        f"В рамках данного исследования добровольцы будут "
        f"госпитализированы в исследовательский центр "
        f"{_fmt_days(hosp_days, '')} вечером накануне приема препаратов "
        f"для проведения процедур {period_names_str} ФК исследования. "
        f"Утром в {_fmt_days(dosing_days)} добровольцы получат однократную "
        f"дозу исследуемого/референтного препарата {intake_cond}. "
        f"Добровольцы останутся в центре в течение как минимум "
        f"{sampling_hours} часов после дозирования (до "
        f"{_fmt_days(discharge_days, 'Дня')} соответственно) с целью "
        f"отбора биообразцов для анализа фармакокинетики и оценки "
        f"параметров безопасности."
    )

    pfk = "Периодах ФК" if n_periods > 1 else "Периоде ФК"
    L.append(
        f"В {pfk} исследования будет проводиться отбор "
        f"{fk_total} проб крови (по {fk_pts} образцов в каждом периоде) "
        f"по 5 мл у каждого добровольца для оценки фармакокинетических "
        f"параметров."
    )

    # График отбора — из blood_sampling расчёта или стандартный
    _sampling_text = None
    if blood_result and blood_result.sampling_times_text:
        _sampling_text = blood_result.sampling_times_text
    elif sampling_times_list:
        # Форматируем из design.sampling_times_hours
        try:
            from app.utils.blood_sampling import _format_sampling_times
        except ImportError:
            try:
                from blood_sampling import _format_sampling_times
            except ImportError:
                _format_sampling_times = None
        if _format_sampling_times:
            _sampling_text = _format_sampling_times(sampling_times_list)

    if _sampling_text:
        L.append(
            f"График отбора образцов крови на ФК: {_sampling_text}."
        )
    else:
        L.append(
            f"График отбора образцов крови на ФК: за 30 мин до приема "
            f"исследуемого препарата/референтного препарата и через "
            f"15, 30, 45 минут, 1, 1.5, 2, 3, 4, 6, 8, 10, 12, 16, 24 "
            f"часов после приема препарата."
        )

    L.append(
        f"Таким образом, выбранная схема отбора проб обеспечивает, "
        f"согласно требованиям п.38 Решения №85 ЕАЭС, надежную "
        f"оценку длительности экспозиции, при соблюдении оптимального "
        f"соотношения риск-польза для добровольцев и отсутствии "
        f"избыточного отбора проб и кровопотери."
    )

    # Блок 6: Объём крови
    L.append(
        f"Общий объем крови для определения фармакокинетических "
        f"параметров {inn} в {_periods_word_gen(n_periods)} для одного "
        f"добровольца составит {fk_total_ml:.1f} мл "
        f"({fk_total} образцов по 5 мл, заполнение системы "
        f"0,5 мл × {flush_count} = {flush_ml:.1f} мл)."
    )

    pd_word = "периода" if n_periods < 5 else "периодов"
    L.append(
        f"Всего из клинического центра в аналитическую лабораторию, "
        f"при условии завершения исследования всеми рандомизированными "
        f"добровольцами, будет передано {biosamples} биообразца плазмы "
        f"({fk_pts} точек × {n_periods} {pd_word} × {n_total} добровольца)."
    )

    L.append(
        f"Объем крови, забираемой для клинического, биохимического "
        f"и серологического анализов, составит {lab_ml:.0f} мл "
        f"за время исследования."
    )

    L.append(
        f"На протяжении всего исследования у каждого добровольца "
        f"будет отобрано {total_ml:.1f} мл крови во всех периодах "
        f"исследования."
    )

    # Блок 7: Отмывочный период
    if n_periods >= 2 and washout_days:
        L.append("Отмывочный период")
        L.append(
            f"Отмывочный период между приемами дозы "
            f"исследуемого/референтного препарата составит "
            f"{washout_days} дней с момента первого приема препаратов "
            f"исследования. Длительность отмывочного периода превышает "
            f"длительность 5-ти периодов полувыведения {inn} "
            f"(Т1/2 {inn} из плазмы крови составляет около "
            f"{t_half_str} часов). Во время отмывочного периода будет "
            f"проводиться оценка параметров безопасности и отбор "
            f"биообразцов крови для фармакокинетического анализа."
        )

    # Блок 8: Наблюдение
    L.append("Период последующего наблюдения")
    L.append(
        f"Визит периода последующего наблюдения будет проведен на "
        f"{follow_up_days} день (День {followup_day}) от последнего "
        f"приема препарата исследования с целью оценки НЯ/СНЯ."
    )
    L.append(
        f"Сбор данных о НЯ/СНЯ и сопутствующей терапии, а также "
        f"оценка жизненно важных показателей будут проводиться во "
        f"время каждого визита добровольца в центр, включая дни "
        f"госпитализации. Информация о НЯ/СНЯ будет также собираться "
        f"во время отмывочного периода и периода последующего наблюдения."
    )

    return "\n".join(L), _rand_table_data


def _methodology_blanks_values(tl, b, n_periods, washout, follow_up,
                                n_total, inn, t_half_hours) -> list:
    """Значения для _replace_blanks_in_runs (шаблон с EN SPACE)."""
    n_per_group = str(n_total // 2)
    fk_pts = str(b.fk_points_per_period) if b else "18"
    fk_total = str(b.fk_total_points) if b else ""
    t_half_str = f"{t_half_hours:.2f}" if t_half_hours and t_half_hours < 1 else (
        f"{t_half_hours:.1f}" if t_half_hours else ""
    )

    dosing_days = tl.dosing_days if tl else [1 + i * washout for i in range(n_periods)]
    discharge_days = [p.sampling_end_day for p in tl.fk_periods] if tl else []
    hosp_days = [p.hospitalization_day for p in tl.fk_periods] if tl else []
    fk_period_days = str(tl.fk_period_days) if tl else "2"

    flush_count = (b.fk_total_points - n_periods) if b else 0
    flush_ml = f"{b.flush_volume_ml:.1f}" if b else ""
    fk_total_ml = f"{b.fk_total_ml:.1f}" if b else ""
    lab_ml = f"{b.lab_total_ml:.0f}" if b else ""
    total_ml = f"{b.total_per_subject_ml:.1f}" if b else ""
    biosamples = str(b.biosamples_total) if b else ""

    # Возвращаем список значений в порядке blanks шаблона (160 blanks)
    # Это приблизительно — для полного шаблона нужно точное соответствие
    return [
        # п.1: скрининг дни, ФК дни, отмывочный, наблюдение
        "14", fk_period_days, str(washout), str(follow_up),
        # п.2-3: группы
        n_per_group, "", "", "", n_per_group, "", "", "",
        # п.7: госпитализация, дни приёма, часы отбора, дни выписки
        str(hosp_days[0]) if hosp_days else "",
        str(dosing_days[0]) if dosing_days else "",
        str(dosing_days[1]) if len(dosing_days) > 1 else "",
        "", "", "", "",
        str(sampling_hours_val) if 'sampling_hours_val' in dir() else "24",
        str(discharge_days[0]) if discharge_days else "",
        str(discharge_days[1]) if len(discharge_days) > 1 else "",
        # п.8: пробы
        fk_total, fk_pts,
        # п.9: график отбора
        "30", "15, 30, 45", "1, 1.5, 2, 3, 4, 6, 8, 10, 12, 16, 24",
        # п.11: объём крови
        inn, fk_total_ml, fk_total, flush_count, flush_ml,
        # п.12: биообразцы
        biosamples, fk_pts, str(n_periods), str(n_total),
        # п.13: лаб анализы
        lab_ml,
        # п.14: итого
        total_ml,
        # п.16: отмывочный
        str(washout), inn, t_half_str, t_half_str,
        # п.18: наблюдение
        str(follow_up), str(dosing_days[-1] + follow_up if dosing_days else ""),
    ]


def _get_anova_factors(design_type: str) -> str:
    """
    Возвращает факторы ANOVA в зависимости от дизайна исследования.

    Решение 85, п.61-63:
    - Перекрёстный: последовательность, период, препарат
    - Репликативный: + добровольца, вложенного в последовательность
    - Параллельный: препарат (группа)
    """
    dt = (design_type or "").lower()

    if "parallel" in dt:
        return "препарат (группа)"
    elif "replicate" in dt or "2x2x4" in dt or "2x2x3" in dt or "4_period" in dt or "3_period" in dt:
        return (
            "последовательность, период, препарат. "
            "Доброволец, вложенный в последовательность, будет включён "
            "как случайный фактор"
        )
    else:
        # Стандартный перекрёстный 2×2
        return "последовательность, период, препарат"


def _insert_hypothesis_formulas(cell):
    """
    Вставляет формулы гипотез БЭ в пустые параграфы P2 и P3 ячейки Row 25.
    H₀₁: G_T/G_R ≤ 0,80  vs  H₁₁: G_T/G_R > 0,80
    H₀₂: G_T/G_R ≥ 1,25  vs  H₁₂: G_T/G_R < 1,25
    """
    from docx.shared import Pt

    def _add(para, text, subscript=False, italic=False, bold=False):
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        if subscript:
            r.font.subscript = True
        if italic:
            r.font.italic = True
        if bold:
            r.font.bold = True
        return r

    # Находим пустые параграфы P2, P3 (формулы)
    empty_paras = []
    for i, p in enumerate(cell.paragraphs):
        if not p.text.strip() and i < 5:
            empty_paras.append(p)

    if len(empty_paras) >= 2:
        # P2: H₀₁: G_T / G_R ≤ 0,80  vs  H₁₁: G_T / G_R > 0,80
        p2 = empty_paras[0]
        _add(p2, "H", italic=True)
        _add(p2, "01", subscript=True)
        _add(p2, ": ", italic=True)
        _add(p2, "G", italic=True)
        _add(p2, "T", subscript=True)
        _add(p2, " / ")
        _add(p2, "G", italic=True)
        _add(p2, "R", subscript=True)
        _add(p2, " ≤ 0,80   ")
        _add(p2, "vs", italic=True)
        _add(p2, "   ")
        _add(p2, "H", italic=True)
        _add(p2, "11", subscript=True)
        _add(p2, ": ", italic=True)
        _add(p2, "G", italic=True)
        _add(p2, "T", subscript=True)
        _add(p2, " / ")
        _add(p2, "G", italic=True)
        _add(p2, "R", subscript=True)
        _add(p2, " > 0,80")

        # P3: H₀₂: G_T / G_R ≥ 1,25  vs  H₁₂: G_T / G_R < 1,25
        p3 = empty_paras[1]
        _add(p3, "H", italic=True)
        _add(p3, "02", subscript=True)
        _add(p3, ": ", italic=True)
        _add(p3, "G", italic=True)
        _add(p3, "T", subscript=True)
        _add(p3, " / ")
        _add(p3, "G", italic=True)
        _add(p3, "R", subscript=True)
        _add(p3, " ≥ 1,25   ")
        _add(p3, "vs", italic=True)
        _add(p3, "   ")
        _add(p3, "H", italic=True)
        _add(p3, "12", subscript=True)
        _add(p3, ": ", italic=True)
        _add(p3, "G", italic=True)
        _add(p3, "T", subscript=True)
        _add(p3, " / ")
        _add(p3, "G", italic=True)
        _add(p3, "R", subscript=True)
        _add(p3, " < 1,25")


def _insert_stat_methods_formatted(cell, inn: str, anova_factors: str):
    """
    Вставляет полный текст Row 25 с форматированием:
    формулы гипотез, AUC_{0-t}, C_{max}, T_{max}, G_T, G_R — subscript.
    """
    from docx.shared import Pt

    # Очищаем ячейку
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    first_para = cell.paragraphs[0]
    for run in first_para.runs:
        run._element.getparent().remove(run._element)

    def _add(para, text, subscript=False, italic=False):
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        if subscript:
            r.font.subscript = True
        if italic:
            r.font.italic = True
        return r

    def _np():
        return cell.add_paragraph()

    # P0: Введение
    p = first_para
    _add(p, "Статистический анализ данных будет проведен в соответствии с подходами, "
            "изложенными в Правилах проведения исследований биоэквивалентности "
            "лекарственных препаратов в рамках Евразийского экономического союза, "
            "при помощи специализированного программного обеспечения, окончательный "
            "выбор которого будет произведен на этапе финализации плана "
            "статистического анализа.")

    # P1: Гипотеза
    p1 = _np()
    _add(p1, "В исследовании будет проверяться гипотеза биоэквивалентности:")

    # P2: H₀₁
    p2 = _np()
    _add(p2, "H", italic=True)
    _add(p2, "01", subscript=True)
    _add(p2, ": ", italic=True)
    _add(p2, "G", italic=True)
    _add(p2, "T", subscript=True)
    _add(p2, " / ")
    _add(p2, "G", italic=True)
    _add(p2, "R", subscript=True)
    _add(p2, " ≤ 0,80   ")
    _add(p2, "vs", italic=True)
    _add(p2, "   ")
    _add(p2, "H", italic=True)
    _add(p2, "11", subscript=True)
    _add(p2, ": ", italic=True)
    _add(p2, "G", italic=True)
    _add(p2, "T", subscript=True)
    _add(p2, " / ")
    _add(p2, "G", italic=True)
    _add(p2, "R", subscript=True)
    _add(p2, " > 0,80")

    # P3: H₀₂
    p3 = _np()
    _add(p3, "H", italic=True)
    _add(p3, "02", subscript=True)
    _add(p3, ": ", italic=True)
    _add(p3, "G", italic=True)
    _add(p3, "T", subscript=True)
    _add(p3, " / ")
    _add(p3, "G", italic=True)
    _add(p3, "R", subscript=True)
    _add(p3, " ≥ 1,25   ")
    _add(p3, "vs", italic=True)
    _add(p3, "   ")
    _add(p3, "H", italic=True)
    _add(p3, "12", subscript=True)
    _add(p3, ": ", italic=True)
    _add(p3, "G", italic=True)
    _add(p3, "T", subscript=True)
    _add(p3, " / ")
    _add(p3, "G", italic=True)
    _add(p3, "R", subscript=True)
    _add(p3, " < 1,25")

    # P4: где GT и GR
    p4 = _np()
    _add(p4, "где ")
    _add(p4, "G", italic=True)
    _add(p4, "T", subscript=True)
    _add(p4, " и ")
    _add(p4, "G", italic=True)
    _add(p4, "R", subscript=True)
    _add(p4, " — средние геометрические ФК параметра исследуемого "
             "препарата и препарата сравнения соответственно.")

    # P5: Описательные статистики
    p5 = _np()
    _add(p5, "Для первичных и вторичных фармакокинетических параметров будут "
             "рассчитаны и приведены описательные статистики (медиана, среднее "
             "арифметическое, среднее геометрическое, минимальное и максимальное "
             "значение, стандартное отклонение, коэффициент вариации).")

    # P6: Логнормальное — с AUC_{0-t}, C_{max}, T_{max}
    p6 = _np()
    _add(p6, "Статистический анализ с целью подтверждения или отклонения гипотезы "
             "о биоэквивалентности будет проведен в предположении о логнормальном "
             "распределении параметров ")
    _add(p6, "AUC")
    _add(p6, "0-t", subscript=True)
    _add(p6, " и ")
    _add(p6, "С")
    _add(p6, "max", subscript=True)
    _add(p6, " и нормальном распределении остальных параметров, за исключением ")
    _add(p6, "Т")
    _add(p6, "max", subscript=True)
    _add(p6, ". Для подтверждения биоэквивалентности будет необходимо достижение "
             "критериев биоэквивалентности, указанных выше, для показателей ")
    _add(p6, "AUC")
    _add(p6, "0-t", subscript=True)
    _add(p6, " и ")
    _add(p6, "С")
    _add(p6, "max", subscript=True)
    _add(p6, f" {inn}.")

    # P7: ANOVA
    p7 = _np()
    _add(p7, "После проведения логарифмического преобразования эти показатели "
             "анализируются с помощью дисперсионного анализа "
             "(analysis of variance, ANOVA).")

    # P8: Факторы
    p8 = _np()
    _add(p8, "Статистическая модель дисперсионного анализа должна включать "
             f"следующие фиксированные факторы: {anova_factors}.")

    # P9: Фактическое время
    p9 = _np()
    _add(p9, "Расчёты будут произведены на основании данных о фактическом, "
             "а не расчётном времени забора каждого образца.")

    # P10: Безопасность
    p10 = _np()
    _add(p10, "Все добровольцы, которые получили хотя бы одну дозу исследуемого "
              "препарата, будут включены в популяцию для анализа безопасности. "
              "Параметры безопасности будут анализироваться при помощи описательной "
              "статистики. Все результаты статистической обработки данных будут "
              "представлены в виде таблиц.")

    # P11: НЯ/MedDRA
    p11 = _np()
    _add(p11, "Все НЯ будут зафиксированы и классифицированы по степени тяжести "
              "и в соответствии с терминологией Медицинского словаря для "
              "регуляторной деятельности (MedDRA) (наиболее поздней версии). "
              "Число НЯ и СНЯ, а также число и процент добровольцев с НЯ и СНЯ "
              "будут представлены в виде таблиц по классу систем органов и "
              "предпочтительному термину, по связи с препаратом и степени тяжести.")

    # P12: Лаборатория
    p12 = _np()
    _add(p12, "Будут проанализированы результаты лабораторных исследований, "
              "включая любые изменения по сравнению с исходными. Отклонения от "
              "нормальных значений будут обобщены в виде частотных таблиц. "
              "Сводные таблицы динамики лабораторных показателей будут "
              "представлены по группам в соответствии с результатами, полученными "
              "на визитах, и сдвигом относительно исходного значения.")

    # P13: Жизненные показатели
    p13 = _np()
    _add(p13, "Будут проанализированы жизненные показатели, данные физикального "
              "осмотра и ЭКГ, включая клинически значимые изменения по сравнению "
              "с исходными. Отклонения от нормы будут обобщены в виде частотных "
              "таблиц.")

    # P14: Непрерывные/дискретные
    p14 = _np()
    _add(p14, "Непрерывные показатели будут представлены средним значением, "
              "стандартным отклонением, медианой, минимальным и максимальным "
              "значением, 95% доверительным интервалом (ДИ). Дискретные "
              "показатели – числом и процентом добровольцев.")


def _generate_blinding_randomization_text(
    n_sequences: int,
    seqs: list,
    n_per_group: str,
    design_type: str = "",
) -> str:
    """
    Генерирует текст Row 26 «Заслепление и Рандомизация».

    Структура:
    1. Заслепление (открытое + лаборатория заслеплена)
    2. Блочная рандомизация (IWRS iRand)
    3. Последовательности
    4. Количество в группе
    """
    n_groups = len(seqs)
    seq_display = "/".join(seqs)  # "TR/RT" или "TRTR/RTRT"

    # Количество групп словом
    groups_word_rod = {2: "двух", 3: "трёх", 4: "четырёх"}.get(n_groups, str(n_groups))

    # Соотношение
    ratio = ":".join(["1"] * n_groups)  # "1:1" или "1:1:1"

    L = []

    # Заслепление
    L.append(
        "Данное исследование является открытым. Однако, для биоаналитической "
        "лаборатории проводится заслепление: сотрудники лаборатории не будут "
        "иметь доступа к рандомизационному списку до окончания биоаналитической "
        "стадии исследования и сдачи таблицы концентраций Спонсору исследования."
    )

    # Рандомизация
    L.append(
        f"Добровольцы будут распределены в одну из {groups_word_rod} групп "
        f"в соответствии с рандомизационным списком методом блочной рандомизации "
        f"без стратификации в соотношении {ratio} с использованием программы "
        f"IWRS iRand. Каждый доброволец будет рандомизирован с присвоением "
        f"рандомизационного номера, который определяет последовательность "
        f"приёма исследуемого препарата (T) и референтного препарата (R) – "
        f"{seq_display}."
    )

    # Количество
    L.append(
        f"В каждую группу будет включено по {n_per_group} добровольцев."
    )

    L.append(
        "В каждом периоде добровольцы получат исследуемый/референтный "
        "препарат по схеме в соответствии с группой, в которую они распределены."
    )

    return "\n".join(L)



def _insert_volunteers_count_formatted(
    cell, inn: str, cv_intra_val, n_with_dropout: str, n_total: str,
    dropout_pct: float, screenfail_pct: float,
):
    """
    Вставляет краткий текст Row 12 «Количество добровольцев»
    с форматированием: CV_{intra}, C_{max}, AUC_{0-t} — subscript.
    """
    from docx.shared import Pt

    # Очищаем ячейку
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    first_para = cell.paragraphs[0]
    for run in first_para.runs:
        run._element.getparent().remove(run._element)

    def _add(para, text, subscript=False, italic=False):
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        if subscript:
            r.font.subscript = True
        if italic:
            r.font.italic = True
        return r

    p = first_para

    # Строка 1: CVintra для Cmax/AUC0-t
    _add(p, "При расчёте объёма выборки для данного исследования "
            "использовался коэффициент внутрииндивидуальной вариабельности "
            "фармакокинетических параметров (")
    _add(p, "CV")
    _add(p, "intra", subscript=True)
    _add(p, f") {cv_intra_val}% для максимальной концентрации/площади "
            "под кривой «концентрация-время» (")
    _add(p, "C")
    _add(p, "max", subscript=True)
    _add(p, "/")
    _add(p, "AUC")
    _add(p, "0-t", subscript=True)
    _add(p, f") {inn}.")

    # Строка 2: R пакет
    p2 = cell.add_paragraph()
    _add(p2, "Расчёт объёма выборки выполнен с помощью пакета ")
    _add(p2, "Power", italic=False)
    _add(p2, "TOST", italic=False)
    _add(p2, " с использованием ПО «The R Project for Statistical Computing» "
             "(https://www.r-project.org) версии не ниже 4.4.2.")

    # Строка 3: dropout
    p3 = cell.add_paragraph()
    _add(p3, f"С учётом досрочного выбывания не более "
             f"{dropout_pct*100:.0f}% включённых добровольцев, "
             f"в исследование будет включено {n_with_dropout} добровольцев.")

    # Строка 4: screen-fail
    p4 = cell.add_paragraph()
    _add(p4, f"С учётом возможного {screenfail_pct*100:.0f}% отсева на "
             f"скрининге в исследование будут скринированы до "
             f"{n_total} добровольцев.")

    # Строка 5: не будут заменены
    p5 = cell.add_paragraph()
    _add(p5, "Добровольцы, досрочно завершившие исследование, "
             "не будут заменены.")


def _insert_be_criteria_formatted(cell, inn: str, be_lower: float, be_upper: float):
    """
    Вставляет критерии БЭ (Row 22) с форматированием:
    AUC_{0-t}, C_{max}, α — subscript/греческие буквы.
    """
    import copy
    from docx.shared import Pt
    from docx.oxml.ns import qn

    # Получаем форматирование
    source_rpr = None
    for para in cell.paragraphs:
        for run in para.runs:
            rpr = run._element.find(qn('w:rPr'))
            if rpr is not None:
                source_rpr = copy.deepcopy(rpr)
                break
        if source_rpr is not None:
            break

    # Очищаем ячейку
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    first_para = cell.paragraphs[0]
    for run in first_para.runs:
        run._element.getparent().remove(run._element)

    def _add(para, text, subscript=False, italic=False):
        r = para.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        if subscript:
            r.font.subscript = True
        if italic:
            r.font.italic = True
        return r

    p = first_para

    _add(p, "Вывод о биоэквивалентности сравниваемых препаратов будет сделан "
            "с использованием подхода, основанного на оценке 90 % доверительных "
            "интервалов для отношений средних геометрических значений "
            "исследуемого препарата (T) к препарату сравнения (R) для "
            "фармакокинетических параметров ")
    _add(p, "AUC")
    _add(p, "0-t", subscript=True)
    _add(p, " и ")
    _add(p, "C")
    _add(p, "max", subscript=True)
    _add(p, f" {inn}. "
            "Препараты считаются биоэквивалентными, если границы оцененных "
            "доверительных интервалов для ")
    _add(p, "AUC")
    _add(p, "0-t", subscript=True)
    _add(p, " и ")
    _add(p, "C")
    _add(p, "max", subscript=True)
    _add(p, f" находятся в пределах {be_lower:.2f}–{be_upper:.2f} % (")
    _add(p, "α", italic=True)
    _add(p, "=0,05) для изучаемого аналита.")


def _insert_pk_parameters_formatted(cell, inn: str):
    """
    Вставляет раздел «Изучаемые ФК параметры» в ячейку Row 20
    с правильным форматированием: subscript для индексов, italic для курсива.

    Формат как в шаблоне:
    C_max, AUC_{0-t}, AUC_{0-∞}, T_{max}, T_{1/2}, k_{el}
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    # Получаем форматирование из ячейки
    source_rpr = None
    for para in cell.paragraphs:
        for run in para.runs:
            rpr = run._element.find(qn('w:rPr'))
            if rpr is not None:
                source_rpr = copy.deepcopy(rpr)
                break
        if source_rpr is not None:
            break

    # Очищаем ячейку
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    first_para = cell.paragraphs[0]
    for run in first_para.runs:
        run._element.getparent().remove(run._element)

    def _apply_fmt(run, subscript=False, italic=False, bold=False):
        """Применяет форматирование к run."""
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        if subscript:
            run.font.subscript = True
        if italic:
            run.font.italic = True
        if bold:
            run.font.bold = True

    def _add_para(text_parts, style=None, bullet=False):
        """
        Добавляет параграф с форматированными частями.
        text_parts: list of (text, {subscript, italic, bold})
        """
        para = cell.add_paragraph()
        if first_para.style:
            para.style = first_para.style
        if bullet:
            # Добавляем маркер списка
            pPr = para._element.get_or_add_pPr()
            numPr = OxmlElement('w:numPr')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            numPr.append(ilvl)
            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), '1')
            numPr.append(numId)
            pPr.append(numPr)

        for text, fmt in text_parts:
            r = para.add_run(text)
            _apply_fmt(
                r,
                subscript=fmt.get("sub", False),
                italic=fmt.get("it", False),
                bold=fmt.get("b", False),
            )
        return para

    N = {}  # normal
    S = {"sub": True}  # subscript
    I = {"it": True}  # italic

    # Используем первый параграф для вводного текста
    r = first_para.add_run(
        f"В данном исследовании будет изучена фармакокинетика {inn} "
        f"по данным исходного соединения в плазме крови."
    )
    _apply_fmt(r)

    # 1. Первичные параметры
    _add_para([
        ("1. Следующие фармакокинетические параметры ", N),
        (f"{inn} ", N),
        ("будут оцениваться в качестве ", N),
        ("первичных", I),
        (":", N),
    ])

    # • C_max
    _add_para([
        ("C", N), ("max", S),
        (" - максимальная плазменная концентрация", N),
    ], bullet=True)

    # • AUC_{0-t}
    _add_para([
        ("AUC", N), ("0-t", S),
        (" - площадь под кривой «плазменная концентрация – время» "
         "с момента приема до последней определяемой концентрации "
         "во временной точке ", N),
        ("t", I),
        (".", N),
    ], bullet=True)

    # 2. Вторичные параметры
    _add_para([
        ("2. Следующие фармакокинетические параметры ", N),
        (f"{inn} ", N),
        ("будут оцениваться в качестве ", N),
        ("вторичных", I),
        (":", N),
    ])

    # • AUC_{0-∞}
    _add_para([
        ("AUC", N), ("0-∞", S),
        (" - площадь под фармакокинетической кривой «концентрация-время», "
         "начиная с нулевого значения времени, экстраполированная "
         "до бесконечности", N),
    ], bullet=True)

    # • T_{max}
    _add_para([
        ("Т", N), ("max", S),
        (" - время достижения максимальной концентрации", N),
    ], bullet=True)

    # • T_{1/2}
    _add_para([
        ("Т", N), ("1/2", S),
        (" - период полувыведения из плазмы", N),
    ], bullet=True)

    # • k_{el}
    _add_para([
        ("k", N), ("el", S),
        (" - константа скорости терминальной элиминации.", N),
    ], bullet=True)


def _generate_periods_text(
    tl,
    n_periods: int,
    washout_days: int,
    follow_up_days: int,
    sampling_hours: int,
    intake_mode: str = "fasting",
    sex_restriction: str = "males_only",
) -> str:
    """
    Генерирует текст для Row 18 «Периоды исследования».

    Структура:
    - Период скрининга
    - Период 1..N ФК исследования (визиты, госпитализация, приём, отбор)
    - Отмывочные периоды между ФК периодами
    - Период последующего наблюдения
    - Незапланированный визит
    - Визит досрочного завершения
    """
    import math

    # Timeline данные
    if tl is not None:
        fk_periods = tl.fk_periods
        dosing_days = tl.dosing_days
        visits = tl.visits
        fk_period_days = tl.fk_period_days
    else:
        # Fallback без timeline
        sampling_cal = math.ceil(sampling_hours / 24) + 1 if sampling_hours > 0 else 2
        fk_period_days = sampling_cal
        dosing_days = [1 + i * washout_days for i in range(n_periods)]
        fk_periods = None

    # Intake текст
    if intake_mode == "fed":
        intake_sentence = (
            "добровольцы получат однократную дозу исследуемого препарата "
            "после приема высококалорийной пищи (прием пищи не ранее, чем за "
            "30 минут до приема исследуемых препаратов), запивая 200 мл "
            "негазированной воды."
        )
    else:
        intake_sentence = (
            "добровольцы получат однократную дозу исследуемого препарата "
            "натощак (после не менее, чем 8-часового ночного голодания до "
            "приема препарата), запивая 200 мл негазированной воды."
        )

    # Тест на беременность
    include_pregnancy = sex_restriction in ("males_and_females", "females_only")
    pregnancy_text = ", тест на беременность" if include_pregnancy else ""

    L = []

    # ── Период скрининга ──
    L.append(
        "Период скрининга (предварительное обследование добровольцев):"
    )
    L.append("Визит 1. (День -14 – День -1).")
    L.append(
        "Для оценки соответствия добровольца критериям включения/невключения "
        "должны быть известны все результаты клинических, лабораторных и иных "
        "обследований добровольца, предусмотренных протоколом в рамках периода "
        "скрининга."
    )

    # ── Периоды ФК ──
    L.append("Периоды ФК исследования:")
    visit_num = 2  # Визит 1 = скрининг

    for p_idx in range(n_periods):
        p_num = p_idx + 1

        if fk_periods is not None:
            fp = fk_periods[p_idx]
            hosp_day = fp.hospitalization_day
            dose_day = fp.dosing_day
            samp_start = fp.sampling_start_day
            samp_end = fp.sampling_end_day
            discharge = fp.discharge_day
        else:
            hosp_day = dosing_days[p_idx] - 1
            dose_day = dosing_days[p_idx]
            samp_start = dose_day
            samp_end = dose_day + fk_period_days - 1
            discharge = samp_end

        L.append(f"Период {p_num} ФК исследования.")
        L.append(
            f"Визит {visit_num}. День {hosp_day} – День {discharge} "
            f"(госпитализация)"
        )

        # День 0 — госпитализация + рандомизация (только для первого периода)
        if p_num == 1:
            L.append(
                f"\tГоспитализация и рандомизация – День {hosp_day}"
            )
            L.append(
                f"На визите День {hosp_day} накануне дозирования будет проведена "
                f"оценка жизненных показателей, физикальный осмотр, тесты на "
                f"наркотики и алкоголь{pregnancy_text}. В День {hosp_day} "
                f"добровольцы будут рандомизированы и будет выдана карточка "
                f"участника исследования."
            )
        else:
            L.append(f"\tГоспитализация – День {hosp_day}")

        L.append(f"\tПрием препарата – День {dose_day}")
        L.append(
            f"В День {dose_day} {intake_sentence}"
        )
        L.append(
            f"\tОтбор образцов крови для анализа фармакокинетики и оценка "
            f"параметров безопасности – День {samp_start} – День {samp_end}"
        )
        L.append(
            f"Добровольцы останутся в центре в течение как минимум "
            f"{sampling_hours} часов после дозирования."
        )
        L.append(
            f"В День {dose_day} для фармакокинетического анализа будет получена "
            f"серия образцов крови до приема исследуемого препарата и после "
            f"приема исследуемого препарата."
        )
        L.append(f"\tЗавершение госпитализации – День {discharge}")

        # Отмывочный период (между ФК периодами, кроме последнего)
        if p_idx < n_periods - 1:
            next_dose = dosing_days[p_idx + 1] if p_idx + 1 < len(dosing_days) else dose_day + washout_days
            wo_start = dose_day + 1
            wo_end = next_dose - 1
            L.append(
                f"Отмывочный период: День {wo_start} – День {wo_end} "
                f"({washout_days} дней от приема препарата в Периоде {p_num} "
                f"ФК исследования)."
            )

        visit_num += 1

    # ── Период последующего наблюдения ──
    last_dose = dosing_days[-1] if dosing_days else 1
    followup_day = last_dose + follow_up_days

    L.append("Период последующего наблюдения:")
    L.append(f"Визит {visit_num}. День {followup_day} (окно визита +2 дня)")
    L.append(
        f"\tДоброволец посетит центр через {follow_up_days} дней с момента "
        f"последнего приема препарата, будет осуществлен сбор данных о "
        f"состоянии добровольца."
    )
    L.append(
        "Сбор данных о НЯ/СНЯ и сопутствующей терапии, а также оценка "
        "жизненно важных показателей будут проводиться во время каждого "
        "визита добровольца в центр, включая дни госпитализации. "
        "Информация о НЯ/СНЯ будет также собираться во время отмывочного "
        "периода и периода последующего наблюдения."
    )

    # ── Незапланированный визит ──
    L.append("Незапланированный визит")
    L.append(
        "Проводится при необходимости. При наличии показаний может быть "
        "дополнительно выполнена любая из процедур исследования по решению "
        "Исследователя."
    )

    # ── Досрочное завершение ──
    L.append("Визит досрочного завершения участия в исследовании")
    L.append(
        "Проводится при досрочном выбывании добровольца из исследования."
    )

    return "\n".join(L)


def _generate_duration_text(
    tl,
    n_periods: int,
    washout_days: int,
    follow_up_days: int,
    sampling_hours: int,
) -> str:
    """
    Генерирует текст для Row 19 «Продолжительность исследования».

    Эталонный формат:
    Максимальная продолжительность участия в исследовании для одного
    добровольца составит 43 дня (не более 7 недель). С учетом окна визита
    наблюдения (+1 день) максимальная длительность исследования может
    составить 44 дня. Период скрининга продлится от 1 до 14 дней
    (День от -14 до -1).
    Длительность Периодов 1, 2, 3 и 4 ФК – 23 дня (от Дня 0 до Дня 22
    включительно), включая 7 дней отмывочного периода в каждом периоде.
    Период наблюдения – 7 дней (День 22 – День 28 включительно) от приема
    последней дозы исследуемого препарата).
    """
    import math

    # Данные из timeline
    if tl is not None:
        total_max = tl.total_days_max
        dosing_days = tl.dosing_days
        last_discharge = tl.fk_periods[-1].discharge_day if tl.fk_periods else 2
    else:
        dosing_days = [1 + i * washout_days for i in range(n_periods)]
        sampling_cal = math.ceil(sampling_hours / 24) + 1 if sampling_hours > 0 else 2
        last_discharge = dosing_days[-1] + sampling_cal - 1
        total_max = 14 + dosing_days[-1] + follow_up_days

    last_dose = dosing_days[-1]
    followup_day = last_dose + follow_up_days

    # ФК часть: от Дня 0 до последнего дня приёма включительно
    fk_total_days = last_dose + 1  # День 0 .. День last_dose = last_dose+1 дней

    # Недели
    total_weeks = math.ceil(total_max / 7)

    # Перечисление периодов
    period_nums = list(range(1, n_periods + 1))
    if len(period_nums) == 1:
        periods_enum = "Периода 1"
    elif len(period_nums) == 2:
        periods_enum = f"Периодов {period_nums[0]} и {period_nums[1]}"
    else:
        periods_enum = (
            "Периодов "
            + ", ".join(str(p) for p in period_nums[:-1])
            + f" и {period_nums[-1]}"
        )

    parts = []

    # Максимальная продолжительность
    parts.append(
        f"Максимальная продолжительность участия в исследовании для одного "
        f"добровольца составит {total_max} дней (не более {total_weeks} недель). "
        f"С учетом окна визита наблюдения (+2 дня) максимальная длительность "
        f"исследования может составить {total_max + 2} дней."
    )

    # Скрининг
    parts.append(
        f"Период скрининга продлится от 1 до 14 дней (День от -14 до -1)."
    )

    # ФК часть
    if n_periods > 1 and washout_days:
        parts.append(
            f"Длительность {periods_enum} ФК – {fk_total_days} дней "
            f"(от Дня 0 до Дня {last_dose} включительно), включая "
            f"{washout_days} дней отмывочного периода в каждом периоде."
        )
    else:
        parts.append(
            f"Длительность {periods_enum} ФК – {fk_total_days} дней "
            f"(от Дня 0 до Дня {last_dose} включительно)."
        )

    # Наблюдение
    # "7 дней от приёма" = День приёма (1-й день) .. День +6 (7-й день)
    followup_end_day = last_dose + follow_up_days - 1
    parts.append(
        f"Период наблюдения – {follow_up_days} дней "
        f"(День {last_dose} – День {followup_end_day} включительно) "
        f"от приема последней дозы исследуемого препарата."
    )

    return " ".join(parts)