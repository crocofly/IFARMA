"""
services/export/rationale_exporter.py — Генератор пояснительного файла.

Создаёт .docx документ с обоснованиями каждого пункта синопсиса:
- Источники ФК-параметров (PubMed, инструкция)
- Обоснование дизайна (дерево решений по Решению №85)
- Расчёт размера выборки (формулы, промежуточные значения)
- Встроенный PK-график (PNG)
- Регуляторные проверки
- Полный список источников
"""

import os
from typing import Any, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─── Утилиты форматирования ───

def _h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
    return h


def _p(doc, text, bold=False, italic=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return para


def _bullet(doc, text, italic=False):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.italic = italic
    return para


def _tbl(doc, headers, rows_data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
    for row_data in rows_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val) if val is not None else "\u2014"
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)
    return table


def export_rationale(
    pipeline_result: Dict[str, Any],
    output_path: str = "rationale.docx",
    pk_curve_path: Optional[str] = None,
    pk_curve_data: Optional[Any] = None,
) -> str:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # -- Распаковка --
    pk = pipeline_result.get("pk", {})
    if hasattr(pk, "model_dump"):
        pk = pk.model_dump()
    design = pipeline_result.get("design", {})
    if hasattr(design, "model_dump"):
        design = design.model_dump()
    sample = pipeline_result.get("sample_size", {})
    if hasattr(sample, "model_dump"):
        sample = sample.model_dump()
    regulatory = pipeline_result.get("regulatory", {})
    synopsis = pipeline_result.get("synopsis", {})

    inn_ru = pk.get("inn_ru") or synopsis.get("inn", "")
    inn_en = pk.get("inn_en", "")
    pk_sources = pk.get("sources", [])
    t_half_hours = pk.get("t_half_hours") or 0
    cv_intra = pk.get("cv_intra") or 0

    # ============ ТИТУЛ ============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "ПОЯСНИТЕЛЬНЫЙ ФАЙЛ\n"
        "К СИНОПСИСУ ПРОТОКОЛА ИССЛЕДОВАНИЯ БИОЭКВИВАЛЕНТНОСТИ"
    )
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(f"Действующее вещество: {inn_ru}")
    r2.font.size = Pt(12)
    r2.font.name = "Times New Roman"
    if inn_en:
        sub2 = doc.add_paragraph()
        sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = sub2.add_run(f"INN: {inn_en}")
        r3.font.size = Pt(12)
        r3.font.name = "Times New Roman"
        r3.italic = True
    doc.add_paragraph()

    # ============ 1. ОБЩАЯ ИНФОРМАЦИЯ ============
    _h(doc, "1. Общая информация о препарате", 1)
    _tbl(doc, ["Параметр", "Значение"], [
        ("МНН (русск.)", inn_ru),
        ("INN (англ.)", inn_en or "\u2014"),
        ("Лекарственная форма", synopsis.get("dosage_form", "\u2014")),
        ("Дозировка", synopsis.get("dosage", "\u2014")),
        ("Референтный препарат", synopsis.get("reference_drug", pk.get("reference_drug", "\u2014"))),
        ("BCS класс", pk.get("bcs_class", "\u2014")),
        ("Высоковариабельный (HVD)", "Да" if pk.get("is_hvd") else "Нет"),
        ("Узкий терап. индекс (NTI)", "Да" if pk.get("is_nti") else "Нет"),
    ])

    # ============ 2. ФК-ПАРАМЕТРЫ ============
    _h(doc, "2. Фармакокинетические параметры", 1)
    _p(doc, (
        "Фармакокинетические параметры извлечены автоматически из открытых "
        "источников: инструкции по медицинскому применению (ГРЛС), публикаций "
        "PubMed/PMC и FDA/EMA review documents. Для каждого параметра указан "
        "первоисточник."
    ))

    # 2.1 Источники
    _h(doc, "2.1. Источники данных", 2)
    if pk_sources:
        for i, src in enumerate(pk_sources, 1):
            src_type = src.get("source_type", "unknown").upper()
            title_text = src.get("title", "\u2014")
            pmid = src.get("pmid", "")
            doi = src.get("doi", "")
            url = src.get("url", "")
            ref_parts = [f"[{i}] [{src_type}] {title_text}"]
            if pmid:
                ref_parts.append(f"PMID: {pmid}")
            if doi:
                ref_parts.append(f"DOI: {doi}")
            if url:
                ref_parts.append(url)
            _bullet(doc, " | ".join(ref_parts))
    else:
        _p(doc, "Источники не найдены.", italic=True)

    # 2.2 Параметры
    _h(doc, "2.2. Извлечённые ФК-параметры", 2)
    pk_keys = [
        ("Cmax", "cmax"), ("Tmax", "tmax"), ("T\u00bd", "t_half"),
        ("AUC\u2080\u208b\u209c", "auc_0t"), ("AUC\u2080\u208b\u221e", "auc_0inf"),
        ("CVintra (Cmax)", "cv_intra_cmax"), ("CVintra (AUC)", "cv_intra_auc"),
    ]
    pk_rows = []
    for label, key in pk_keys:
        param = pk.get(key)
        if param and isinstance(param, dict):
            pk_rows.append((label, f"{param.get('value', '\u2014')} {param.get('unit', '')}", param.get("source", "\u2014")))
        elif param is not None and not isinstance(param, dict):
            pk_rows.append((label, str(param), "\u2014"))
    if pk_rows:
        _tbl(doc, ["Параметр", "Значение", "Источник"], pk_rows)

    _p(doc, "")
    _p(doc, "Ключевые параметры для выбора дизайна:", bold=True)
    _bullet(doc, f"T\u00bd = {t_half_hours} ч \u2014 определяет возможность перекрёстного дизайна и длительность отмывочного периода")
    _bullet(doc, f"CVintra = {cv_intra}% \u2014 определяет необходимость репликативного дизайна (порог 30% по Решению \u211685)")
    if pk.get("is_hvd"):
        _bullet(doc, "CVintra \u2265 30% \u2192 препарат классифицирован как высоковариабельный (HVD)")
    if pk.get("is_nti"):
        _bullet(doc, "Препарат классифицирован как NTI (узкий терапевтический индекс)")

    lit = pk.get("literature_review", "")
    if lit:
        _h(doc, "2.3. Литературный обзор", 2)
        _p(doc, lit)

    # ============ 3. PK-ПРОФИЛЬ ============
    if pk_curve_path or pk_curve_data:
        _h(doc, "3. Теоретический фармакокинетический профиль", 1)
        _p(doc, (
            "На основании литературных данных (Cmax, Tmax, T\u00bd) построена "
            "теоретическая кривая \u00abконцентрация \u2014 время\u00bb по однокомпартментной "
            "модели с абсорбцией первого порядка (функция Бейтмана)."
        ))
        _p(doc, "Кривая используется для:", bold=True)
        _bullet(doc, "Обоснования схемы и количества точек отбора проб крови")
        _bullet(doc, "Визуализации ожидаемого фармакокинетического профиля")
        _bullet(doc, "Оценки теоретических значений AUC\u2080\u208b\u209c и AUC\u2080\u208b\u221e")
        _bullet(doc, "Проверки достаточности длительности отбора проб (остаточная AUC < 20%)")

        if pk_curve_path and os.path.exists(pk_curve_path):
            _p(doc, "")
            doc.add_picture(pk_curve_path, width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(f"Рисунок 1. Теоретический фармакокинетический профиль {inn_ru}")
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
            r.italic = True

        if pk_curve_data and hasattr(pk_curve_data, "cmax"):
            cd = pk_curve_data
            _p(doc, "")
            _p(doc, "Параметры модели Бейтмана:", bold=True)
            _tbl(doc, ["Параметр", "Значение"], [
                ("Модель", "Однокомпартментная, абсорбция 1-го порядка"),
                ("Cmax", f"{cd.cmax:.1f} нг/мл"),
                ("Tmax", f"{cd.tmax:.2f} ч"),
                ("T\u00bd", f"{cd.t_half:.2f} ч"),
                ("kel", f"{cd.kel:.4f} ч\u207b\u00b9"),
                ("ka", f"{cd.ka:.4f} ч\u207b\u00b9"),
                ("AUC\u2080\u208b\u209c (теоретич.)", f"{cd.auc_0t:.1f} нг\u00b7ч/мл"),
                ("AUC\u2080\u208b\u221e (теоретич.)", f"{cd.auc_0inf:.1f} нг\u00b7ч/мл"),
                ("Остаточная AUC", f"{cd.auc_residual_pct:.1f}%"),
            ])
            _p(doc, "")
            if cd.auc_residual_pct < 20:
                _p(doc, f"\u2705 Остаточная AUC = {cd.auc_residual_pct:.1f}% < 20%. "
                        "Длительность отбора проб достаточна (AUC\u2080\u208b\u209c \u2265 80% от AUC\u2080\u208b\u221e).")
            else:
                _p(doc, f"\u26a0\ufe0f Остаточная AUC = {cd.auc_residual_pct:.1f}% \u2265 20%. "
                        "Рекомендуется увеличить длительность отбора.")
            _p(doc, "Основание: Решение ЕЭК \u211685, п. 38.", italic=True)

            if hasattr(cd, "sampling_times") and cd.sampling_times:
                _p(doc, "")
                _p(doc, "Предложенная схема отбора крови:", bold=True)
                def _fmt(t):
                    if t == 0: return "0 (предозовая)"
                    if t < 1: return f"{t*60:.0f} мин"
                    if t == int(t): return f"{int(t)} ч"
                    return f"{t:.2f} ч"
                _bullet(doc, f"Количество точек: {len(cd.sampling_times)}")
                _bullet(doc, f"Времена: {', '.join(_fmt(t) for t in cd.sampling_times)}")

    # ============ 4. ОБОСНОВАНИЕ ДИЗАЙНА ============
    _h(doc, "4. Обоснование дизайна исследования", 1)

    d_type = design.get("design_type", "")
    labels = {
        "2x2_crossover": "Простой перекрёстный 2\u00d72 (TR/RT)",
        "replicate_3_period": "Частично репликативный 3-периодный",
        "replicate_4_period": "Полный репликативный 4-периодный (TRTR/RTRT)",
        "full_replicate_4_period": "Полный репликативный 4-периодный (TRTR/RTRT)",
        "full_replicate_3_period": "Полный репликативный 3-периодный",
        "parallel": "Параллельный (2 группы)",
        "parallel_1_period": "Параллельный однопериодный",
        "adaptive_crossover": "Адаптивный перекрёстный (двухэтапный)",
        "adaptive_parallel_1": "Адаптивный параллельный (двухэтапный)",
    }
    _p(doc, f"Выбранный дизайн: {labels.get(d_type, d_type)}", bold=True)
    _p(doc, "")
    _p(doc, "Дерево решений (Решение ЕЭК \u211685):", bold=True)

    if "parallel" in (d_type or ""):
        _bullet(doc, f"T\u00bd = {t_half_hours} ч \u2192 отмывочный \u2265 {7*t_half_hours/24:.0f} дней \u2192 неприемлемо")
        _bullet(doc, "Выбор: параллельный дизайн (п. 25 Решения \u211685)")
    elif pk.get("is_nti"):
        _bullet(doc, "Препарат NTI \u2192 репликативный дизайн")
        _bullet(doc, "Границы: 90.00\u2013111.11% (п. 50 Решения \u211685)")
    elif pk.get("is_hvd"):
        _bullet(doc, f"CVintra = {cv_intra}% \u2265 30% \u2192 высоковариабельный")
        _bullet(doc, "Выбор: репликативный дизайн + ABEL (п. 48\u201349 Решения \u211685)")
    else:
        _bullet(doc, f"CVintra = {cv_intra}% < 30%, T\u00bd = {t_half_hours} ч \u2192 стандартный")
        _bullet(doc, "Выбор: перекрёстный 2\u00d72 (п. 21\u201324 Решения \u211685)")

    just = design.get("design_justification", "")
    if just:
        _p(doc, "")
        _p(doc, "Развёрнутое обоснование:", bold=True)
        _p(doc, just)

    n_periods = design.get("n_periods", 2)
    washout = design.get("washout_days", 0)
    _p(doc, "")
    _p(doc, "Параметры дизайна:", bold=True)
    intake_map = {"fasting": "Натощак", "fed": "После еды", "both": "Натощак и после еды"}
    _tbl(doc, ["Параметр", "Значение"], [
        ("Тип дизайна", labels.get(d_type, d_type)),
        ("Периодов", str(n_periods)),
        ("Последовательностей", str(design.get("n_sequences", 2))),
        ("Последовательности", design.get("sequences_description", "\u2014")),
        ("Отмывочный период", f"{washout} дней"),
        ("Формула отмывочного", design.get("washout_formula", "") or f"\u2265 5 \u00d7 T\u00bd = {5*t_half_hours:.1f} ч \u2192 {washout} дней"),
        ("Режим приёма", intake_map.get(design.get("intake_mode", "fasting"), "\u2014")),
        ("Точки отбора крови", str(design.get("n_blood_points", "\u2014"))),
        ("Длительность отбора", f"{design.get('sampling_duration_hours', '\u2014')} ч"),
        ("Follow-up", f"{design.get('follow_up_days', 7)} дней"),
    ])

    # Границы БЭ
    _p(doc, "")
    _p(doc, "Границы биоэквивалентности:", bold=True)
    be_method = design.get("be_method", "standard")
    if be_method == "ABEL":
        _tbl(doc, ["Параметр", "Значение"], [
            ("Метод", "ABEL (scaled Average BE with Expanding Limits)"),
            ("AUC\u2080\u208b\u209c", f"{design.get('be_lower',80):.2f}\u2013{design.get('be_upper',125):.2f}%"),
            ("Cmax (расширенные)", f"{design.get('be_lower_cmax',80):.2f}\u2013{design.get('be_upper_cmax',125):.2f}%"),
            ("PE constraints", f"{design.get('be_pe_lower',80):.2f}\u2013{design.get('be_pe_upper',125):.2f}%"),
            ("Регуляторная константа", str(design.get("be_regulatory_constant", 0.76))),
        ])
        _p(doc, "Основание: Решение ЕЭК \u211685, п. 48\u201349.", italic=True)
    elif be_method == "NTID":
        _tbl(doc, ["Параметр", "Значение"], [
            ("Метод", "NTID"), ("Границы", "90.00\u2013111.11%"),
        ])
        _p(doc, "Основание: Решение ЕЭК \u211685, п. 50.", italic=True)
    else:
        _tbl(doc, ["Параметр", "Значение"], [
            ("Метод", "Стандартный ABE (TOST)"),
            ("Границы", f"{design.get('be_lower',80):.2f}\u2013{design.get('be_upper',125):.2f}%"),
        ])
        _p(doc, "Основание: Решение ЕЭК \u211685, п. 44\u201347.", italic=True)

    # ============ 5. РАСЧЁТ ВЫБОРКИ ============
    _h(doc, "5. Расчёт размера выборки", 1)

    cv_used = sample.get("cv_intra_used", cv_intra)
    n_base = sample.get("n_base", "\u2014")
    n_dropout = sample.get("n_with_dropout", "\u2014")
    n_total = sample.get("n_total", "\u2014")
    dropout_rate = sample.get("dropout_rate", 0.15)
    screenfail_rate = sample.get("screenfail_rate", 0.15)
    power = sample.get("power", 0.8)
    gmr = sample.get("gmr", 0.95)

    _p(doc, "Метод расчёта:", bold=True)
    if pk.get("is_hvd") and n_periods >= 4:
        _p(doc, f"sampleN.scABEL.ad (R/PowerTOST) \u2014 симуляционный метод для ABEL, "
                f"CVintra = {cv_used}%, дизайн {n_periods}-периодный.")
    elif pk.get("is_nti"):
        _p(doc, f"sampleN.NTID (R/PowerTOST) \u2014 границы \u03b8\u2080 = 0.975.")
    else:
        _p(doc, f"sampleN.TOST (R/PowerTOST) \u2014 стандартный TOST, CVintra = {cv_used}%.")

    _p(doc, "")
    _p(doc, "Входные параметры:", bold=True)
    _tbl(doc, ["Параметр", "Значение"], [
        ("CVintra", f"{cv_used}%"),
        ("Мощность (1\u2212\u03b2)", f"{power*100:.0f}%"),
        ("\u03b1", str(sample.get("alpha", 0.05))),
        ("GMR (\u03b8\u2080)", str(gmr)),
        ("Дизайн", f"{n_periods}-периодный"),
    ])

    _p(doc, "")
    _p(doc, "Результат:", bold=True)
    _tbl(doc, ["Параметр", "Значение"], [
        ("n базовый (завершившие)", str(n_base)),
        (f"n рандомизированных (+{dropout_rate*100:.0f}% dropout)", str(n_dropout)),
        (f"n скринированных (+{screenfail_rate*100:.0f}% screen-fail)", str(n_total)),
    ])

    _p(doc, "")
    _p(doc, "Пояснение:", bold=True)
    _bullet(doc, f"n_base = {n_base} (таблица PowerTOST, CVintra={cv_used}%, мощность={power*100:.0f}%, GMR={gmr})")
    _bullet(doc, f"n_рандом = \u2308{n_base} / (1 \u2212 {dropout_rate})\u2309 = {n_dropout}")
    _bullet(doc, f"n_скрининг = \u2308{n_dropout} / (1 \u2212 {screenfail_rate})\u2309 = {n_total}")

    calc_desc = sample.get("calculation_description", "")
    if calc_desc:
        _p(doc, "")
        _p(doc, "Вывод R/PowerTOST:", bold=True)
        _p(doc, calc_desc, size=9)

    _p(doc, "")
    _p(doc, "Основание: Решение ЕЭК \u211685, п. 27 (минимум 12 добровольцев); "
            "Приложение 5 (стат. анализ). R/PowerTOST v1.5-6.", italic=True)

    blood = sample.get("blood_volume_ml", 0)
    if blood:
        _p(doc, "")
        _p(doc, "Контроль объёма крови:", bold=True)
        _bullet(doc, f"Общий объём: {blood:.0f} мл")
        _bullet(doc, "\u2705 \u2264 450 мл" if blood <= 450 else "\u274c > 450 мл!")

    # ============ 6. РЕГУЛЯТОРНАЯ ПРОВЕРКА ============
    _h(doc, "6. Регуляторная проверка", 1)
    _p(doc, "Автоматическая проверка на соответствие Решению ЕЭК \u211685 и Дополнению \u211630.")

    summary = regulatory.get("summary", {})
    if summary:
        _p(doc, f"Вердикт: {summary.get('verdict', '\u2014')}", bold=True)
        _tbl(doc, ["Статус", "Кол-во"], [
            ("\u2705 PASS", str(summary.get("pass", 0))),
            ("\u274c FAIL", str(summary.get("fail", 0))),
            ("\u26a0\ufe0f WARNING", str(summary.get("warning", 0))),
            ("\u2753 MISSING", str(summary.get("missing_data", 0))),
        ])

    checks = regulatory.get("checks", [])
    if checks:
        _p(doc, "")
        _p(doc, "Детальные проверки:", bold=True)
        section = None
        for ch in checks:
            s = ch.get("section", "")
            if s != section:
                section = s
                _p(doc, f"\u2014 {s} \u2014", bold=True)
            icon = {
                "PASS": "\u2705", "FAIL": "\u274c",
                "WARNING": "\u26a0\ufe0f", "MISSING_DATA": "\u2753", "NA": "\u2796",
            }.get(ch.get("status", ""), "?")
            _bullet(doc, f"{icon} [{ch.get('id','')}] {ch.get('rule','')}")
            det = ch.get("detail", "")
            if det:
                _p(doc, f"     {det}", italic=True, size=10)
            ref = ch.get("reference", "")
            if ref:
                _p(doc, f"     Основание: {ref}", italic=True, size=9)

    # ============ 7. ИСТОЧНИКИ ============
    _h(doc, "7. Использованные источники", 1)

    _p(doc, "Литературные источники:", bold=True)
    seen = set()
    all_src = []
    for src in pk_sources:
        t = src.get("title", "")
        if t and t not in seen:
            seen.add(t)
            parts = [t]
            if src.get("pmid"): parts.append(f"PMID: {src['pmid']}")
            if src.get("doi"): parts.append(f"DOI: {src['doi']}")
            if src.get("url"): parts.append(src["url"])
            all_src.append(" | ".join(parts))
    for s in pipeline_result.get("sources", []):
        if isinstance(s, str) and s not in seen:
            seen.add(s)
            all_src.append(s)
    if all_src:
        for i, s in enumerate(all_src, 1):
            _bullet(doc, f"[{i}] {s}")
    else:
        _p(doc, "Не указаны.", italic=True)

    _p(doc, "")
    _p(doc, "Нормативные документы:", bold=True)
    for n in [
        "Решение Совета ЕЭК от 03.11.2016 \u211685 \u00abПравила проведения исследований "
        "биоэквивалентности лекарственных препаратов в рамках ЕАЭС\u00bb",
        "Решение Совета ЕЭК от 03.11.2016 \u211679 \u00abПравила Надлежащей клинической практики ЕАЭС\u00bb",
        "Дополнение \u211630 от 12.04.2024 к Решению \u211685",
        "ГОСТ Р 57679-2017 \u00abЛекарственные средства. Исследование биоэквивалентности\u00bb",
        "Хельсинкская декларация ВМА (75-я Ген. Ассамблея, октябрь 2024 г.)",
    ]:
        _bullet(doc, n)

    _p(doc, "")
    _p(doc, "Программное обеспечение:", bold=True)
    _bullet(doc, "R / PowerTOST v1.5-6 (расчёт размера выборки)")
    _bullet(doc, "IWRS iRand (рандомизация)")

    # ============ DISCLAIMER ============
    doc.add_page_break()
    _h(doc, "Отказ от ответственности", 2)
    _p(doc, (
        "Настоящий документ сформирован автоматически мультиагентной системой "
        "на основе открытых источников и нормативных документов. Все выводы "
        "требуют проверки квалифицированным специалистом. Система не несёт "
        "ответственности за принятые на основе данного документа решения."
    ), italic=True)

    doc.save(output_path)
    return output_path